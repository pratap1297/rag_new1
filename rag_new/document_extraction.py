import os
import sys
import json
import time
from typing import Optional, List, Dict, Any, Generator, Tuple
from pathlib import Path
from enum import Enum
from abc import ABC, abstractmethod
from datetime import datetime
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
import logging
import base64
import mimetypes
from collections import Counter, defaultdict
import re

# Azure imports
from azure.ai.inference import ChatCompletionsClient, EmbeddingsClient
from azure.ai.inference.models import SystemMessage, UserMessage
from azure.core.credentials import AzureKeyCredential
from azure.cognitiveservices.vision.computervision import ComputerVisionClient
from azure.cognitiveservices.vision.computervision.models import (
    VisualFeatureTypes, OperationStatusCodes, ReadOperationResult
)
from msrest.authentication import CognitiveServicesCredentials

# Document processing imports
try:
    import pdfplumber
    from docx import Document as DocxDocument
    import openpyxl
    from PIL import Image
    import fitz  # PyMuPDF for better PDF handling
    import easyocr  # Fallback OCR
    import pytesseract  # Another fallback OCR
    from langdetect import detect
    import textstat
    from dotenv import load_dotenv
except ImportError as e:
    print(f"Missing required packages: {e}")
    print("Install with: pip install pdfplumber python-docx openpyxl pillow PyMuPDF easyocr pytesseract langdetect textstat azure-ai-inference azure-cognitiveservices-vision-computervision python-dotenv")
    sys.exit(1)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class BaseAdapter(ABC):
    """Base adapter for document extraction."""
    
    supported_mimetypes: List[str] = []
    needs_vision: bool = False
    needs_llm: bool = False
    
    @abstractmethod
    def extract(self, file_path: str, options: Dict[str, Any]) -> Generator['ExtractionUnit', None, None]:
        """Extract content from file."""
        pass


class ExtractionMethod(Enum):
    NATIVE = "native"
    VISION_READ = "vision_read"
    VISION_OCR = "vision_ocr"
    EASYOCR = "easyocr"
    TESSERACT = "tesseract"
    LLAMA4 = "llama4"
    MIXED = "mixed"


class ContentType(Enum):
    TEXT = "text"
    TABLE = "table"
    FIGURE_CAPTION = "figure_caption"
    HEADING = "heading"
    LIST = "list"
    FORM_FIELD = "form_field"
    HANDWRITTEN = "handwritten"
    MIXED = "mixed"


class DocumentIntelligence:
    """Intelligent document analysis and metadata extraction."""
    
    def __init__(self, llm_client: Optional[ChatCompletionsClient] = None, 
                 llm_model: Optional[str] = None):
        self.llm_client = llm_client
        self.llm_model = llm_model
    
    def analyze_content(self, text: str) -> Dict[str, Any]:
        """Analyze content to determine type, quality, and metadata."""
        analysis = {
            'content_type': self._detect_content_type(text),
            'language': self._detect_language(text),
            'readability_score': textstat.flesch_reading_ease(text),
            'complexity': self._assess_complexity(text),
            'key_entities': self._extract_entities(text),
            'quality_score': self._assess_quality(text),
            'needs_review': False
        }
        
        # Flag for review if quality is low
        if analysis['quality_score'] < 0.7:
            analysis['needs_review'] = True
        
        return analysis
    
    def generate_summary(self, text: str, max_length: int = 200) -> str:
        """Generate intelligent summary of text."""
        if not self.llm_client or not text.strip():
            return self._basic_summary(text, max_length)
        
        try:
            messages = [
                {"role": "system", "content": "You are a document analysis expert. Create concise, informative summaries."},
                {"role": "user", "content": f"Summarize this text in max {max_length} characters:\n\n{text[:3000]}"}
            ]
            
            response = self.llm_client.complete(
                messages=messages,
                model=self.llm_model,
                temperature=0.3,
                max_tokens=100
            )
            
            if response.choices:
                return response.choices[0].message.content.strip()
        except Exception as e:
            logger.warning(f"LLM summary failed: {e}")
        
        return self._basic_summary(text, max_length)
    
    def _detect_content_type(self, text: str) -> str:
        """Detect the type of content."""
        if self._is_table(text):
            return ContentType.TABLE.value
        elif self._is_list(text):
            return ContentType.LIST.value
        elif self._is_form(text):
            return ContentType.FORM_FIELD.value
        elif len(text.split()) < 10 and text.isupper():
            return ContentType.HEADING.value
        else:
            return ContentType.TEXT.value
    
    def _detect_language(self, text: str) -> str:
        """Detect language of text."""
        try:
            return detect(text)
        except:
            return "en"
    
    def _assess_complexity(self, text: str) -> str:
        """Assess text complexity."""
        score = textstat.flesch_reading_ease(text)
        if score >= 90:
            return "very_easy"
        elif score >= 60:
            return "easy"
        elif score >= 30:
            return "medium"
        elif score >= 0:
            return "difficult"
        else:
            return "very_difficult"
    
    def _extract_entities(self, text: str) -> List[str]:
        """Extract key entities from text."""
        # Simple entity extraction - in production, use NLP libraries
        entities = []
        
        # Extract emails
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        entities.extend([f"email:{e}" for e in emails[:5]])
        
        # Extract phone numbers
        phones = re.findall(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text)
        entities.extend([f"phone:{p}" for p in phones[:5]])
        
        # Extract dates
        dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b', text)
        entities.extend([f"date:{d}" for d in dates[:5]])
        
        return entities
    
    def _assess_quality(self, text: str) -> float:
        """Assess extraction quality."""
        if not text.strip():
            return 0.0
        
        score = 1.0
        
        # Check for garbage characters
        garbage_ratio = len(re.findall(r'[^\x00-\x7F]+', text)) / len(text)
        score -= garbage_ratio * 0.5
        
        # Check for reasonable word length
        words = text.split()
        if words:
            avg_word_length = sum(len(w) for w in words) / len(words)
            if avg_word_length > 15 or avg_word_length < 2:
                score -= 0.3
        
        # Check for repetitive patterns
        if len(set(words)) < len(words) * 0.3:
            score -= 0.3
        
        return max(0.0, min(1.0, score))
    
    def _basic_summary(self, text: str, max_length: int) -> str:
        """Create basic summary without LLM."""
        sentences = text.split('.')
        summary = ""
        for sentence in sentences:
            if len(summary) + len(sentence) < max_length:
                summary += sentence + "."
            else:
                break
        return summary.strip()
    
    def _is_table(self, text: str) -> bool:
        """Check if text represents tabular data."""
        lines = text.split('\n')
        if len(lines) < 2:
            return False
        
        # Check for consistent delimiters
        tab_counts = [line.count('\t') for line in lines[:5]]
        pipe_counts = [line.count('|') for line in lines[:5]]
        
        return (all(c > 1 for c in tab_counts) or all(c > 1 for c in pipe_counts))
    
    def _is_list(self, text: str) -> bool:
        """Check if text is a list."""
        lines = text.split('\n')
        list_patterns = [r'^\d+\.', r'^-', r'^•', r'^\*']
        matches = sum(1 for line in lines if any(re.match(p, line.strip()) for p in list_patterns))
        return matches > len(lines) * 0.5
    
    def _is_form(self, text: str) -> bool:
        """Check if text contains form fields."""
        form_indicators = ['name:', 'date:', 'address:', 'phone:', 'email:', 'signature:']
        text_lower = text.lower()
        matches = sum(1 for indicator in form_indicators if indicator in text_lower)
        return matches >= 3


class ExtractionUnit:
    """Enhanced extraction unit with summary and rich metadata."""
    
    def __init__(self, text: str, metadata: Dict[str, Any], page_num: int = 1):
        self.text = text
        self.page_num = page_num
        self.metadata = metadata
        self.metadata['hash'] = {'sha256': hashlib.sha256(text.encode()).hexdigest()}
        self.metadata['extraction']['timestamp_utc'] = datetime.utcnow().isoformat()
        self.metadata['source']['page'] = page_num
        
        # Add text statistics
        self.metadata['statistics'] = {
            'char_count': len(text),
            'word_count': len(text.split()),
            'line_count': len(text.split('\n'))
        }


class WordAdapter(BaseAdapter):
    """Adapter for Word documents."""
    
    supported_mimetypes = ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]
    
    def __init__(self):
        self.doc_intelligence = DocumentIntelligence()
    
    def extract(self, file_path: str, options: Dict[str, Any]) -> Generator[ExtractionUnit, None, None]:
        """Extract text from Word documents."""
        try:
            doc = DocxDocument(file_path)
            
            # Combine all paragraphs for better context
            full_text = []
            paragraph_count = 0
            
            for para_num, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
                    paragraph_count += 1
            
            # Process tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_text.append(row_text)
                if table_text:
                    full_text.append('\n'.join(table_text))
            
            # Extract as single unit with all content
            if full_text:
                combined_text = '\n\n'.join(full_text)
                
                # Analyze content
                content_analysis = self.doc_intelligence.analyze_content(combined_text)
                summary = self.doc_intelligence.generate_summary(combined_text, max_length=200)
                
                metadata = {
                    'source': {
                        'file_path': file_path,
                        'adapter': 'word_adapter',
                        'page': 1,
                        'total_paragraphs': paragraph_count,
                        'block_id': 'full-document'
                    },
                    'extraction': {
                        'method': ExtractionMethod.NATIVE.value,
                        'confidence': 1.0
                    },
                    'content': {
                        'type': content_analysis['content_type'],
                        'language': content_analysis['language'],
                        'summary': summary,
                        'quality_score': content_analysis['quality_score']
                    },
                    'entities': content_analysis['key_entities'],
                    'pipeline': {'version': '2.0.0'}
                }
                
                yield ExtractionUnit(combined_text, metadata, 1)
                
        except Exception as e:
            logger.error(f"Error extracting Word document {file_path}: {e}")


class ExcelAdapter(BaseAdapter):
    """Enhanced Excel adapter with image extraction support."""
    
    supported_mimetypes = ["application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", 
                         "application/vnd.ms-excel"]
    
    def __init__(self, fallback_extractor: Optional['UniversalFallbackExtractor'] = None):
        self.fallback_extractor = fallback_extractor
        self.doc_intelligence = DocumentIntelligence()
    
    def extract(self, file_path: str, options: Dict[str, Any]) -> Generator[ExtractionUnit, None, None]:
        """Extract data and images from Excel files."""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            for sheet_idx, sheet_name in enumerate(workbook.sheetnames):
                sheet = workbook[sheet_name]
                
                # Extract text content
                content = []
                has_data = False
                
                # Get all cell values
                for row in sheet.iter_rows():
                    row_data = []
                    for cell in row:
                        if cell.value is not None:
                            has_data = True
                            row_data.append(str(cell.value))
                        else:
                            row_data.append('')
                    
                    if any(row_data):  # Only include non-empty rows
                        content.append('\t'.join(row_data))
                
                # Extract charts/images if present
                if hasattr(sheet, '_images') and sheet._images:
                    for img_idx, image in enumerate(sheet._images):
                        # Save image temporarily
                        temp_img_path = f"temp_excel_img_{sheet_idx}_{img_idx}.png"
                        try:
                            with open(temp_img_path, 'wb') as f:
                                f.write(image._data())
                            
                            # Extract text from image if it contains any
                            if self.fallback_extractor:
                                img_text, confidence, method = self.fallback_extractor.extract_with_fallback(
                                    temp_img_path,
                                    [ExtractionMethod.VISION_OCR, ExtractionMethod.EASYOCR]
                                )
                                if img_text.strip():
                                    content.append(f"\n[Embedded Image {img_idx+1}]:\n{img_text}")
                        except Exception as e:
                            logger.warning(f"Failed to extract image from Excel: {e}")
                        finally:
                            if os.path.exists(temp_img_path):
                                os.remove(temp_img_path)
                
                if has_data and content:
                    text = '\n'.join(content)
                    
                    # Analyze content
                    content_analysis = self.doc_intelligence.analyze_content(text)
                    summary = self.doc_intelligence.generate_summary(text, max_length=200)
                    
                    # Detect if it's a form, table, or mixed content
                    content_type = ContentType.TABLE.value
                    if self._is_form_like(content):
                        content_type = ContentType.FORM_FIELD.value
                    
                    metadata = {
                        'source': {
                            'file_path': file_path,
                            'adapter': 'excel_adapter',
                            'page': sheet_idx + 1,
                            'sheet_name': sheet_name,
                            'total_sheets': len(workbook.sheetnames),
                            'block_id': f'sheet-{sheet_name}'
                        },
                        'extraction': {
                            'method': ExtractionMethod.NATIVE.value,
                            'confidence': 1.0
                        },
                        'content': {
                            'type': content_type,
                            'language': content_analysis['language'],
                            'summary': summary,
                            'has_images': hasattr(sheet, '_images') and len(sheet._images) > 0,
                            'row_count': sheet.max_row,
                            'column_count': sheet.max_column
                        },
                        'entities': content_analysis['key_entities'],
                        'pipeline': {'version': '2.0.0'}
                    }
                    
                    yield ExtractionUnit(text, metadata, sheet_idx + 1)
                    
        except Exception as e:
            logger.error(f"Error extracting Excel file {file_path}: {e}")
    
    def _is_form_like(self, content: List[str]) -> bool:
        """Detect if Excel sheet contains form-like data."""
        # Look for key-value patterns
        form_patterns = 0
        for line in content[:20]:  # Check first 20 rows
            if ':' in line or any(indicator in line.lower() 
                                for indicator in ['name', 'date', 'address', 'phone', 'email']):
                form_patterns += 1
        
        return form_patterns > 5


class ImageAdapter(BaseAdapter):
    """Adapter for image files."""
    
    supported_mimetypes = ["image/jpeg", "image/png", "image/tiff", "image/bmp"]
    needs_vision = True
    
    def __init__(self, vision_client: Optional[ComputerVisionClient] = None, 
                 llm_client: Optional[ChatCompletionsClient] = None,
                 llm_model: Optional[str] = None):
        self.vision_client = vision_client
        self.llm_client = llm_client
        self.llm_model = llm_model
        self.fallback_extractor = None  # Will be set later to avoid circular dependency
        self.doc_intelligence = DocumentIntelligence(llm_client, llm_model)
    
    def set_fallback_extractor(self, fallback_extractor):
        """Set fallback extractor after initialization to avoid circular dependency."""
        self.fallback_extractor = fallback_extractor
    
    def extract(self, file_path: str, options: Dict[str, Any]) -> Generator[ExtractionUnit, None, None]:
        """Extract text from images using Computer Vision or LLM."""
        if not self.fallback_extractor:
            logger.error("Fallback extractor not set for ImageAdapter")
            return
            
        confidence_threshold = options.get('llm_threshold', 0.85)
        
        # Define extraction preference
        preference = [
            ExtractionMethod.VISION_READ,
            ExtractionMethod.VISION_OCR,
            ExtractionMethod.LLAMA4,
            ExtractionMethod.EASYOCR,
            ExtractionMethod.TESSERACT
        ]
        
        # Extract text using fallback methods
        text, confidence, method = self.fallback_extractor.extract_with_fallback(file_path, preference)
        
        if text.strip():
            # Analyze content
            content_analysis = self.doc_intelligence.analyze_content(text)
            summary = self.doc_intelligence.generate_summary(text, max_length=200)
            
            metadata = {
                'source': {
                    'file_path': file_path,
                    'adapter': 'image_adapter',
                    'page': 1,
                    'block_id': 'image-1'
                },
                'extraction': {
                    'method': method.value,
                    'confidence': confidence
                },
                'content': {
                    'type': content_analysis['content_type'],
                    'language': content_analysis['language'],
                    'summary': summary,
                    'quality_score': content_analysis['quality_score']
                },
                'entities': content_analysis['key_entities'],
                'pipeline': {'version': '2.0.0'}
            }
            
            yield ExtractionUnit(text, metadata, 1)


class UniversalFallbackExtractor:
    """Universal fallback extractor that handles any document type."""
    
    def __init__(self, vision_client: Optional[ComputerVisionClient] = None,
                 llm_client: Optional[ChatCompletionsClient] = None,
                 llm_model: Optional[str] = None):
        self.vision_client = vision_client
        self.llm_client = llm_client
        self.llm_model = llm_model
        self.doc_intelligence = DocumentIntelligence(llm_client, llm_model)
        
        # Initialize fallback OCR engines
        try:
            self.easyocr_reader = easyocr.Reader(['en'])
        except:
            self.easyocr_reader = None
            logger.warning("EasyOCR not available")
    
    def extract_with_fallback(self, file_path: str, method_preference: List[ExtractionMethod]) -> Tuple[str, float, ExtractionMethod]:
        """Extract text using multiple fallback methods."""
        text = ""
        confidence = 0.0
        method_used = ExtractionMethod.NATIVE
        
        for method in method_preference:
            try:
                if method == ExtractionMethod.VISION_READ and self.vision_client:
                    text, confidence = self._extract_with_vision_read(file_path)
                elif method == ExtractionMethod.VISION_OCR and self.vision_client:
                    text, confidence = self._extract_with_vision_ocr(file_path)
                elif method == ExtractionMethod.EASYOCR and self.easyocr_reader:
                    text, confidence = self._extract_with_easyocr(file_path)
                elif method == ExtractionMethod.TESSERACT:
                    text, confidence = self._extract_with_tesseract(file_path)
                elif method == ExtractionMethod.LLAMA4 and self.llm_client:
                    text, confidence = self._extract_with_llm(file_path)
                
                if text and confidence > 0.5:
                    method_used = method
                    break
                    
            except Exception as e:
                logger.warning(f"Method {method.value} failed: {e}")
                continue
        
        # If all methods fail, try to extract any readable content
        if not text:
            text = self._last_resort_extraction(file_path)
            confidence = 0.3
            method_used = ExtractionMethod.MIXED
        
        return text, confidence, method_used
    
    def _extract_with_vision_read(self, file_path: str) -> Tuple[str, float]:
        """Use Azure Computer Vision Read API for complex documents."""
        with open(file_path, 'rb') as image_stream:
            read_response = self.vision_client.read_in_stream(
                image_stream,
                raw=True
            )
        
        # Get operation ID from headers
        operation_location = read_response.headers["Operation-Location"]
        operation_id = operation_location.split("/")[-1]
        
        # Wait for completion
        while True:
            read_result = self.vision_client.get_read_result(operation_id)
            if read_result.status not in ['notStarted', 'running']:
                break
            time.sleep(1)
        
        text_blocks = []
        total_confidence = 0
        line_count = 0
        
        if read_result.status == OperationStatusCodes.succeeded:
            for page in read_result.analyze_result.read_results:
                for line in page.lines:
                    text_blocks.append(line.text)
                    # Calculate average confidence from words
                    if hasattr(line, 'words'):
                        for word in line.words:
                            if hasattr(word, 'confidence'):
                                total_confidence += word.confidence
                                line_count += 1
        
        avg_confidence = total_confidence / line_count if line_count > 0 else 0.9
        return '\n'.join(text_blocks), avg_confidence
    
    def _extract_with_vision_ocr(self, file_path: str) -> Tuple[str, float]:
        """Use Azure Computer Vision OCR API."""
        with open(file_path, 'rb') as image_stream:
            ocr_result = self.vision_client.recognize_printed_text_in_stream(
                image_stream,
                detect_orientation=True
            )
        
        text_blocks = []
        for region in ocr_result.regions:
            for line in region.lines:
                line_text = ' '.join([word.text for word in line.words])
                text_blocks.append(line_text)
        
        return '\n'.join(text_blocks), 0.85
    
    def _extract_with_easyocr(self, file_path: str) -> Tuple[str, float]:
        """Use EasyOCR as fallback."""
        result = self.easyocr_reader.readtext(file_path)
        text_blocks = []
        confidences = []
        
        for (bbox, text, confidence) in result:
            text_blocks.append(text)
            confidences.append(confidence)
        
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0.7
        return '\n'.join(text_blocks), avg_confidence
    
    def _extract_with_tesseract(self, file_path: str) -> Tuple[str, float]:
        """Use Tesseract as fallback."""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            # Tesseract doesn't provide confidence directly, estimate based on output
            confidence = 0.7 if text.strip() else 0.0
            return text, confidence
        except Exception as e:
            logger.error(f"Tesseract failed: {e}")
            return "", 0.0
    
    def _extract_with_llm(self, file_path: str) -> Tuple[str, float]:
        """Use multimodal LLM for complex extraction."""
        try:
            # Convert image to base64
            with open(file_path, 'rb') as image_file:
                image_base64 = base64.b64encode(image_file.read()).decode('utf-8')
            
            messages = [
                {
                    "role": "system", 
                    "content": "You are an expert at extracting text from images. Extract ALL text exactly as it appears, preserving formatting and structure."
                },
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract all text from this image. Preserve the exact formatting."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
                    ]
                }
            ]
            
            response = self.llm_client.complete(
                messages=messages,
                model=self.llm_model,
                temperature=0,
                max_tokens=4000
            )
            
            if response.choices:
                return response.choices[0].message.content, 0.95
                
        except Exception as e:
            logger.error(f"LLM extraction failed: {e}")
        
        return "", 0.0
    
    def _last_resort_extraction(self, file_path: str) -> str:
        """Last resort extraction - try to get any readable content."""
        try:
            # Try to read as binary and decode
            with open(file_path, 'rb') as f:
                content = f.read()
                
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
                try:
                    text = content.decode(encoding, errors='ignore')
                    # Filter out non-printable characters
                    text = ''.join(c for c in text if c.isprintable() or c.isspace())
                    if len(text) > 50:  # If we got meaningful content
                        return text
                except:
                    continue
                    
        except Exception as e:
            logger.error(f"Last resort extraction failed: {e}")
        
        return ""


class EnhancedPDFAdapter(BaseAdapter):
    """Enhanced PDF adapter with page-wise extraction and fallback."""
    
    supported_mimetypes = ["application/pdf"]
    
    def __init__(self, fallback_extractor: Optional[UniversalFallbackExtractor] = None):
        self.fallback_extractor = fallback_extractor
        self.doc_intelligence = DocumentIntelligence()
        self.header_footer_cache = {}
    
    def extract(self, file_path: str, options: Dict[str, Any]) -> Generator[ExtractionUnit, None, None]:
        """Extract text from PDF page by page with fallback."""
        try:
            # Try with PyMuPDF first (better for complex PDFs)
            pdf_document = fitz.open(file_path)
            
            # First pass: detect headers and footers
            if options.get('remove_headers_footers', True):
                self._detect_headers_footers(pdf_document)
            
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text = page.get_text()
                confidence = 1.0
                method = ExtractionMethod.NATIVE
                
                # Remove detected headers/footers
                if options.get('remove_headers_footers', True) and text.strip():
                    text = self._remove_headers_footers(text, page_num)
                
                # If native extraction fails or returns garbage, use fallback
                if not text.strip() or _is_garbage_text(text):
                    if self.fallback_extractor:
                        # Convert page to image for OCR
                        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))  # 2x scale for better OCR
                        img_data = pix.pil_tobytes(format="PNG")
                        temp_img_path = f"temp_page_{page_num}.png"
                        
                        with open(temp_img_path, 'wb') as f:
                            f.write(img_data)
                        
                        try:
                            text, confidence, method = self.fallback_extractor.extract_with_fallback(
                                temp_img_path,
                                [ExtractionMethod.VISION_READ, ExtractionMethod.EASYOCR, ExtractionMethod.TESSERACT]
                            )
                            # Remove headers/footers from OCR text too
                            if options.get('remove_headers_footers', True) and text.strip():
                                text = self._remove_headers_footers(text, page_num)
                        finally:
                            if os.path.exists(temp_img_path):
                                os.remove(temp_img_path)
                
                if text.strip():
                    # Analyze content
                    content_analysis = self.doc_intelligence.analyze_content(text)
                    
                    # Generate summary
                    summary = self.doc_intelligence.generate_summary(text, max_length=200)
                    
                    metadata = {
                        'source': {
                            'file_path': file_path,
                            'adapter': 'enhanced_pdf_adapter',
                            'page': page_num + 1,
                            'total_pages': len(pdf_document),
                            'block_id': f'page-{page_num + 1}'
                        },
                        'extraction': {
                            'method': method.value,
                            'confidence': confidence,
                            'fallback_used': method != ExtractionMethod.NATIVE
                        },
                        'content': {
                            'type': content_analysis['content_type'],
                            'language': content_analysis['language'],
                            'summary': summary,
                            'readability_score': content_analysis['readability_score'],
                            'complexity': content_analysis['complexity'],
                            'quality_score': content_analysis['quality_score']
                        },
                        'entities': content_analysis['key_entities'],
                        'flags': ['needs_review'] if content_analysis['needs_review'] else [],
                        'pipeline': {'version': '2.0.0'}
                    }
                    
                    yield ExtractionUnit(text, metadata, page_num + 1)
            
            pdf_document.close()
            
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            # Try pdfplumber as final fallback
            yield from self._extract_with_pdfplumber(file_path)
    
    def _extract_with_pdfplumber(self, file_path: str):
        """Fallback to pdfplumber."""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text and text.strip():
                        content_analysis = self.doc_intelligence.analyze_content(text)
                        summary = self.doc_intelligence.generate_summary(text, max_length=200)
                        
                        metadata = {
                            'source': {
                                'file_path': file_path,
                                'adapter': 'enhanced_pdf_adapter',
                                'page': page_num,
                                'block_id': f'page-{page_num}'
                            },
                            'extraction': {
                                'method': ExtractionMethod.NATIVE.value,
                                'confidence': 0.9,
                                'fallback_used': True
                            },
                            'content': {
                                'type': content_analysis['content_type'],
                                'language': content_analysis['language'],
                                'summary': summary
                            },
                            'pipeline': {'version': '2.0.0'}
                        }
                        yield ExtractionUnit(text, metadata, page_num)
        except Exception as e:
            logger.error(f"PDFPlumber also failed: {e}")
    
    def _detect_headers_footers(self, pdf_document):
        """Detect repeated headers and footers across pages."""
        if len(pdf_document) < 3:  # Need at least 3 pages to detect patterns
            return
        
        # Sample first, middle, and last pages
        sample_pages = [0, len(pdf_document)//2, len(pdf_document)-1]
        page_texts = []
        
        for page_num in sample_pages:
            page = pdf_document[page_num]
            text = page.get_text()
            lines = text.strip().split('\n')
            page_texts.append(lines)
        
        # Find common headers (first 5 lines)
        potential_headers = []
        for i in range(min(5, min(len(pt) for pt in page_texts))):
            lines = [pt[i] for pt in page_texts if i < len(pt)]
            if len(set(lines)) == 1 and lines[0].strip():  # All same and not empty
                potential_headers.append(lines[0])
            else:
                break  # Stop at first non-matching line
        
        # Find common footers (last 5 lines)
        potential_footers = []
        for i in range(1, min(6, min(len(pt) for pt in page_texts))):
            lines = [pt[-i] for pt in page_texts if len(pt) >= i]
            if len(set(lines)) == 1 and lines[0].strip():  # All same and not empty
                potential_footers.append(lines[0])
            else:
                break
        
        self.header_footer_cache['headers'] = potential_headers
        self.header_footer_cache['footers'] = list(reversed(potential_footers))
    
    def _remove_headers_footers(self, text: str, page_num: int) -> str:
        """Remove detected headers and footers from text."""
        if not self.header_footer_cache:
            return text
        
        lines = text.strip().split('\n')
        
        # Remove headers
        headers = self.header_footer_cache.get('headers', [])
        if headers and len(lines) > len(headers):
            # Check if the beginning matches
            if all(lines[i].strip() == headers[i].strip() 
                   for i in range(len(headers)) 
                   if i < len(lines)):
                lines = lines[len(headers):]
        
        # Remove footers
        footers = self.header_footer_cache.get('footers', [])
        if footers and len(lines) > len(footers):
            # Check if the end matches
            if all(lines[-(i+1)].strip() == footers[-(i+1)].strip() 
                   for i in range(len(footers)) 
                   if i < len(lines)):
                lines = lines[:-len(footers)]
        
        # Also remove page numbers if they appear alone
        if lines and re.match(r'^\d+$', lines[-1].strip()):
            lines = lines[:-1]  # Remove last line if it's a page number
        
        return '\n'.join(lines)


class UniversalAdapter(BaseAdapter):
    """Universal adapter that can handle any file type."""
    
    supported_mimetypes = ["*/*"]
    needs_vision = True
    needs_llm = True
    
    def __init__(self, fallback_extractor: UniversalFallbackExtractor):
        self.fallback_extractor = fallback_extractor
        self.doc_intelligence = DocumentIntelligence()
    
    def extract(self, file_path: str, options: Dict[str, Any]) -> Generator[ExtractionUnit, None, None]:
        """Extract from any file type using universal fallback."""
        mime_type, _ = mimetypes.guess_type(file_path)
        
        # For image-like files or unknown types, convert to image
        if mime_type and mime_type.startswith('image'):
            yield from self._extract_from_image(file_path, options)
        else:
            # Try to convert to images and extract
            yield from self._extract_from_unknown(file_path, options)
    
    def _extract_from_image(self, file_path: str, options: Dict[str, Any]):
        """Extract from image files."""
        preference = [
            ExtractionMethod.VISION_READ,
            ExtractionMethod.VISION_OCR,
            ExtractionMethod.LLAMA4,
            ExtractionMethod.EASYOCR,
            ExtractionMethod.TESSERACT
        ]
        
        text, confidence, method = self.fallback_extractor.extract_with_fallback(file_path, preference)
        
        if text.strip():
            content_analysis = self.doc_intelligence.analyze_content(text)
            summary = self.doc_intelligence.generate_summary(text, max_length=200)
            
            metadata = {
                'source': {
                    'file_path': file_path,
                    'adapter': 'universal_adapter',
                    'page': 1,
                    'block_id': 'full-image'
                },
                'extraction': {
                    'method': method.value,
                    'confidence': confidence
                },
                'content': {
                    'type': content_analysis['content_type'],
                    'language': content_analysis['language'],
                    'summary': summary,
                    'quality_score': content_analysis['quality_score']
                },
                'entities': content_analysis['key_entities'],
                'flags': ['universal_extraction'],
                'pipeline': {'version': '2.0.0'}
            }
            
            yield ExtractionUnit(text, metadata, 1)
    
    def _extract_from_unknown(self, file_path: str, options: Dict[str, Any]):
        """Extract from unknown file types."""
        # This is where you'd implement conversion logic
        # For now, try direct text extraction
        text = self.fallback_extractor._last_resort_extraction(file_path)
        
        if text.strip():
            content_analysis = self.doc_intelligence.analyze_content(text)
            summary = self.doc_intelligence.generate_summary(text, max_length=200)
            
            metadata = {
                'source': {
                    'file_path': file_path,
                    'adapter': 'universal_adapter',
                    'page': 1,
                    'block_id': 'full-document'
                },
                'extraction': {
                    'method': ExtractionMethod.MIXED.value,
                    'confidence': 0.5
                },
                'content': {
                    'type': content_analysis['content_type'],
                    'language': content_analysis['language'],
                    'summary': summary
                },
                'flags': ['low_confidence', 'needs_review'],
                'pipeline': {'version': '2.0.0'}
            }
            
            yield ExtractionUnit(text, metadata, 1)


class EnhancedDocumentExtractionPipeline:
    """Enhanced pipeline with universal fallback support."""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.vision_client = self._init_vision_client()
        self.llm_client = self._init_llm_client()
        self.embeddings_client = self._init_embeddings_client()
        
        # Initialize fallback extractor
        self.fallback_extractor = UniversalFallbackExtractor(
            self.vision_client,
            self.llm_client,
            config.get('CHAT_MODEL')
        )
        
        # Initialize document intelligence
        self.doc_intelligence = DocumentIntelligence(
            self.llm_client,
            config.get('CHAT_MODEL')
        )
        
        self.adapters = self._initialize_adapters()
    
    def _initialize_adapters(self) -> Dict[str, BaseAdapter]:
        """Initialize all available adapters with fallback support."""
        adapters = {}
        
        # Enhanced adapters with fallback
        adapters['pdf'] = EnhancedPDFAdapter(self.fallback_extractor)
        adapters['word'] = WordAdapter()
        adapters['excel'] = ExcelAdapter(self.fallback_extractor)
        
        # Initialize image adapter and set fallback extractor
        image_adapter = ImageAdapter(self.vision_client, self.llm_client, self.config.get('CHAT_MODEL'))
        image_adapter.set_fallback_extractor(self.fallback_extractor)
        adapters['image'] = image_adapter
        
        adapters['universal'] = UniversalAdapter(self.fallback_extractor)
        
        return adapters
    
    def _init_vision_client(self) -> Optional[ComputerVisionClient]:
        """Initialize Computer Vision client."""
        endpoint = self.config.get('COMPUTER_VISION_ENDPOINT')
        key = self.config.get('COMPUTER_VISION_KEY')
        
        if endpoint and key:
            credentials = CognitiveServicesCredentials(key)
            return ComputerVisionClient(endpoint, credentials)
        return None
    
    def _init_llm_client(self) -> Optional[ChatCompletionsClient]:
        """Initialize LLM client."""
        endpoint = self.config.get('AZURE_CHAT_ENDPOINT')
        key = self.config.get('AZURE_API_KEY')
        
        if endpoint and key:
            return ChatCompletionsClient(
                endpoint=endpoint,
                credential=AzureKeyCredential(key),
                api_version="2024-05-01-preview"
            )
        return None
    
    def _init_embeddings_client(self) -> Optional[EmbeddingsClient]:
        """Initialize embeddings client."""
        endpoint = self.config.get('AZURE_EMBEDDINGS_ENDPOINT')
        key = self.config.get('AZURE_API_KEY')
        
        if endpoint and key:
            return EmbeddingsClient(
                endpoint=endpoint,
                credential=AzureKeyCredential(key)
            )
        return None
    
    def process_document(self, file_path: str) -> List[ExtractionUnit]:
        """Process a single document with intelligent extraction."""
        ext = Path(file_path).suffix.lower()
        adapter_map = {
            '.pdf': 'pdf',
            '.docx': 'word',
            '.xlsx': 'excel',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.png': 'image',
            '.tiff': 'image',
            '.bmp': 'image'
        }
        
        # Use specific adapter if available, otherwise universal
        adapter_name = adapter_map.get(ext, 'universal')
        adapter = self.adapters[adapter_name]
        
        options = {
            'llm_threshold': self.config.get('LLM_THRESHOLD', 0.85),
            'enable_summary': True,
            'enable_metadata_enrichment': True,
            'remove_headers_footers': True
        }
        
        units = []
        try:
            for unit in adapter.extract(file_path, options):
                # Enrich with embeddings if available
                if self.embeddings_client and self.config.get('EMBEDDING_MODEL'):
                    try:
                        response = self.embeddings_client.embed(
                            input=[unit.text[:8000]],  # Limit text length for embedding
                            model=self.config['EMBEDDING_MODEL']
                        )
                        if response.data:
                            unit.metadata['embedding'] = {
                                'model': self.config['EMBEDDING_MODEL'],
                                'dimensions': len(response.data[0].embedding),
                                'preview': response.data[0].embedding[:5]
                            }
                    except Exception as e:
                        logger.warning(f"Failed to generate embedding: {e}")
                
                units.append(unit)
                
        except Exception as e:
            logger.error(f"Adapter failed for {file_path}: {e}")
            # Try universal adapter as last resort
            if adapter_name != 'universal':
                logger.info("Falling back to universal adapter")
                for unit in self.adapters['universal'].extract(file_path, options):
                    units.append(unit)
        
        return units
    
    def process_folder(self, folder_path: str, output_path: str, 
                      max_depth: int = 5, max_workers: int = 4):
        """Process all documents in a folder with enhanced extraction."""
        results = []
        
        # Discover files
        files = self._discover_files(folder_path, max_depth)
        logger.info(f"Found {len(files)} files to process")
        
        # Process files in parallel
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_file = {
                executor.submit(self.process_document, file_path): file_path 
                for file_path in files
            }
            
            for future in as_completed(future_to_file):
                file_path = future_to_file[future]
                try:
                    units = future.result()
                    results.extend(units)
                    logger.info(f"Processed {file_path}: {len(units)} pages/units extracted")
                    
                    # Log summary for each unit
                    for unit in units:
                        summary = unit.metadata.get('content', {}).get('summary', 'No summary')
                        logger.info(f"  Page {unit.page_num}: {summary[:100]}...")
                        
                except Exception as e:
                    logger.error(f"Failed to process {file_path}: {e}")
        
        # Generate document-level summaries
        doc_summaries = self._generate_document_summaries(results)
        
        # Write results
        self._write_results(results, output_path, doc_summaries)
        logger.info(f"Extraction complete. {len(results)} units written to {output_path}")
    
    def _discover_files(self, folder_path: str, max_depth: int) -> List[str]:
        """Discover all files in folder."""
        files = []
        
        for root, _, filenames in os.walk(folder_path):
            depth = root.replace(folder_path, '').count(os.sep)
            if depth > max_depth:
                continue
                
            for filename in filenames:
                # Process any file type
                if not filename.startswith('.'):  # Skip hidden files
                    files.append(os.path.join(root, filename))
        
        return files
    
    def _generate_document_summaries(self, units: List[ExtractionUnit]) -> Dict[str, Dict[str, Any]]:
        """Generate document-level summaries and insights."""
        doc_summaries = {}
        
        # Group units by file path
        file_units = {}
        for unit in units:
            file_path = unit.metadata['source']['file_path']
            if file_path not in file_units:
                file_units[file_path] = []
            file_units[file_path].append(unit)
        
        # Generate summary for each document
        for file_path, doc_units in file_units.items():
            # Combine text from all pages
            full_text = '\n\n'.join([unit.text[:1000] for unit in doc_units])  # Limit for summary
            
            # Generate comprehensive summary
            if self.llm_client and self.config.get('CHAT_MODEL'):
                try:
                    messages = [
                        {
                            "role": "system",
                            "content": "You are a document analysis expert. Provide comprehensive document summaries with key insights."
                        },
                        {
                            "role": "user",
                            "content": f"""Analyze this document and provide:
                            1. Executive summary (2-3 sentences)
                            2. Document type and purpose
                            3. Key topics covered
                            4. Notable entities or data points
                            5. Quality assessment
                            
                            Document text (first few pages):
                            {full_text[:3000]}"""
                        }
                    ]
                    
                    response = self.llm_client.complete(
                        messages=messages,
                        model=self.config['CHAT_MODEL'],
                        temperature=0.3,
                        max_tokens=500
                    )
                    
                    if response.choices:
                        analysis = response.choices[0].message.content
                        
                        doc_summaries[file_path] = {
                            'llm_analysis': analysis,
                            'basic_summary': f"Document with {len(doc_units)} pages",
                            'page_count': len(doc_units),
                            'total_words': sum(unit.metadata['statistics']['word_count'] for unit in doc_units),
                            'extraction_methods': list(set(unit.metadata['extraction']['method'] for unit in doc_units)),
                            'languages': list(set(unit.metadata['content'].get('language', 'unknown') for unit in doc_units)),
                            'average_quality': sum(unit.metadata['content'].get('quality_score', 0.5) for unit in doc_units) / len(doc_units)
                        }
                except Exception as e:
                    logger.error(f"LLM summary generation failed: {e}")
                        
            # Fallback summary if LLM fails
            if file_path not in doc_summaries:
                doc_summaries[file_path] = {
                    'basic_summary': f"Document with {len(doc_units)} pages",
                    'page_count': len(doc_units),
                    'total_words': sum(unit.metadata['statistics']['word_count'] for unit in doc_units),
                    'extraction_methods': list(set(unit.metadata['extraction']['method'] for unit in doc_units)),
                    'languages': list(set(unit.metadata['content'].get('language', 'unknown') for unit in doc_units)),
                    'average_quality': sum(unit.metadata['content'].get('quality_score', 0.5) for unit in doc_units) / len(doc_units),
                    'llm_analysis': f"Document with {len(doc_units)} pages"
                }
        
        return doc_summaries
        
    def _write_results(self, results: List[ExtractionUnit], output_path: str, doc_summaries: Dict[str, Dict[str, Any]]):
        """Write extraction results with summaries."""
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        
        # Write detailed extraction results
        with open(output_path, 'w', encoding='utf-8') as f:
            for unit in results:
                record = {
                    'text': unit.text,
                    'metadata': unit.metadata,
                    'page_num': unit.page_num
                }
                f.write(json.dumps(record, ensure_ascii=False) + '\n')
        
        # Write document summaries
        summary_path = output_path.replace('.jsonl', '_summaries.json')
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(doc_summaries, f, ensure_ascii=False, indent=2)
        
        # Write extraction report
        report_path = output_path.replace('.jsonl', '_report.txt')
        with open(report_path, 'w', encoding='utf-8') as f:
            f.write("DOCUMENT EXTRACTION REPORT\n")
            f.write("=" * 80 + "\n\n")
            f.write(f"Total documents processed: {len(doc_summaries)}\n")
            f.write(f"Total pages/units extracted: {len(results)}\n\n")
            
            for file_path, summary in doc_summaries.items():
                f.write(f"\nDocument: {os.path.basename(file_path)}\n")
                f.write("-" * 40 + "\n")
                if 'llm_analysis' in summary:
                    f.write(f"Analysis:\n{summary['llm_analysis']}\n")
                else:
                    f.write(f"Summary: {summary.get('basic_summary', 'N/A')}\n")
                f.write(f"Pages: {summary.get('page_count', 'N/A')}\n")
                f.write(f"Total words: {summary.get('total_words', 'N/A')}\n")
                quality_score = summary.get('average_quality', 'N/A')
                if isinstance(quality_score, (int, float)):
                    f.write(f"Quality score: {quality_score:.2f}\n")
                else:
                    f.write(f"Quality score: {quality_score}\n")
                f.write("\n")


def load_config() -> Dict[str, Any]:
    """Load configuration from environment."""
    load_dotenv()
    
    return {
        'AZURE_API_KEY': os.getenv('AZURE_API_KEY'),
        'AZURE_CHAT_ENDPOINT': os.getenv('AZURE_CHAT_ENDPOINT'),
        'AZURE_EMBEDDINGS_ENDPOINT': os.getenv('AZURE_EMBEDDINGS_ENDPOINT'),
        'COMPUTER_VISION_KEY': os.getenv('COMPUTER_VISION_KEY'),
        'COMPUTER_VISION_ENDPOINT': os.getenv('COMPUTER_VISION_ENDPOINT'),
        'CHAT_MODEL': os.getenv('CHAT_MODEL'),
        'EMBEDDING_MODEL': os.getenv('EMBEDDING_MODEL'),
        'LLM_THRESHOLD': float(os.getenv('LLM_THRESHOLD', '0.85'))
    }


def _is_garbage_text(text: str) -> bool:
    """Check if extracted text is garbage."""
    if not text.strip():
        return True
        
    # Check for too many special characters
    special_char_ratio = len(re.findall(r'[^\w\s]', text)) / len(text)
    if special_char_ratio > 0.5:
        return True
        
    # Check for reasonable word structure
    words = text.split()
    if not words:
        return True
            
    avg_word_length = sum(len(w) for w in words) / len(words)
    if avg_word_length > 20 or avg_word_length < 2:
        return True
        
    return False


def main():
    """Main entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Enhanced Document Extraction Pipeline')
    parser.add_argument('--input', '-i', required=True, help='Input folder path or single file')
    parser.add_argument('--output', '-o', default='output/extractions.jsonl', 
                       help='Output file path')
    parser.add_argument('--max-depth', type=int, default=5, 
                       help='Maximum folder depth to traverse')
    parser.add_argument('--workers', type=int, default=4, 
                       help='Number of parallel workers')
    parser.add_argument('--test-fallback', action='store_true',
                       help='Test fallback extraction on a single file')
    
    args = parser.parse_args()
    
    # Load configuration
    config = load_config()
    
    # Initialize pipeline
    pipeline = EnhancedDocumentExtractionPipeline(config)
    
    if args.test_fallback and os.path.isfile(args.input):
        # Test mode for single file
        logger.info(f"Testing extraction on: {args.input}")
        units = pipeline.process_document(args.input)
        
        for unit in units:
            print(f"\n{'='*80}")
            print(f"Page {unit.page_num}:")
            print(f"Method: {unit.metadata['extraction']['method']}")
            print(f"Confidence: {unit.metadata['extraction']['confidence']:.2f}")
            print(f"Language: {unit.metadata['content'].get('language', 'unknown')}")
            print(f"Summary: {unit.metadata['content'].get('summary', 'No summary')}")
            print(f"Text preview: {unit.text[:200]}...")
            print(f"Quality score: {unit.metadata['content'].get('quality_score', 'N/A')}")
    else:
        # Normal processing mode
        if os.path.isfile(args.input):
            # Single file processing
            units = pipeline.process_document(args.input)
            # Generate document summaries
            doc_summaries = pipeline._generate_document_summaries(units)
            pipeline._write_results(units, args.output, doc_summaries)
        else:
            # Folder processing
            pipeline.process_folder(
                args.input,
                args.output,
                max_depth=args.max_depth,
                max_workers=args.workers
            )


if __name__ == "__main__":
    main()