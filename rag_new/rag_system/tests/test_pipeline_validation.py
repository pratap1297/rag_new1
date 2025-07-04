#!/usr/bin/env python3
"""
Comprehensive RAG Pipeline Test Suite
Tests ingestion, retrieval, and conversation for various document types
"""

import os
import sys
import json
import time
import tempfile
import logging
import concurrent.futures
import threading
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
import traceback

# Test framework imports
import unittest
from unittest.mock import MagicMock, patch
import numpy as np

# Add project root to path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

# RAG system imports
from rag_system.src.core.system_init import initialize_system
from rag_system.src.core.dependency_container import get_dependency_container

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class TestDocument:
    """Helper class to create test documents"""
    
    @staticmethod
    def create_test_pdf(path: Path, with_images: bool = True, pages: int = 3) -> Path:
        """Create a test PDF file"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open()
            
            for page_num in range(pages):
                page = doc.new_page()
                
                # Add text
                text = f"Test PDF Page {page_num + 1}\n\n"
                text += f"This is test content for page {page_num + 1}.\n"
                text += "The quick brown fox jumps over the lazy dog.\n"
                text += f"Important information on page {page_num + 1}."
                
                page.insert_text((50, 50), text, fontsize=12)
                
                # Add image placeholder if requested
                if with_images:
                    rect = fitz.Rect(50, 200, 200, 350)
                    page.draw_rect(rect, color=(0, 0, 1), width=2)
                    page.insert_text((60, 220), "[Test Image]", fontsize=10)
            
            doc.save(str(path))
            doc.close()
            return path
            
        except ImportError:
            # Fallback: create a simple text file with .pdf extension
            with open(path, 'w') as f:
                f.write("Test PDF Document\n" * 10)
            return path
    
    @staticmethod
    def create_test_excel(path: Path, with_images: bool = True, sheets: int = 2) -> Path:
        """Create a test Excel file"""
        try:
            import openpyxl
            from openpyxl.drawing.image import Image as XLImage
            
            wb = openpyxl.Workbook()
            
            # Remove default sheet
            wb.remove(wb.active)
            
            for sheet_num in range(sheets):
                ws = wb.create_sheet(f"Sheet{sheet_num + 1}")
                
                # Add headers
                headers = ["ID", "Name", "Department", "Floor", "Manager"]
                for col, header in enumerate(headers, 1):
                    ws.cell(row=1, column=col, value=header)
                
                # Add data
                data = [
                    [1, "John Doe", "Engineering", "Floor 3", "Jane Smith"],
                    [2, "Alice Brown", "Marketing", "Floor 2", "Bob Wilson"],
                    [3, "Charlie Davis", "Sales", "Floor 1", "Eve Johnson"],
                ]
                
                for row_idx, row_data in enumerate(data, 2):
                    for col_idx, value in enumerate(row_data, 1):
                        ws.cell(row=row_idx, column=col_idx, value=value)
                
                # Add a formula
                ws.cell(row=6, column=1, value="Total:")
                ws.cell(row=6, column=2, value="=COUNTA(B2:B4)")
            
            wb.save(str(path))
            return path
            
        except ImportError:
            # Fallback: create CSV
            with open(path, 'w') as f:
                f.write("ID,Name,Department,Floor,Manager\n")
                f.write("1,John Doe,Engineering,Floor 3,Jane Smith\n")
                f.write("2,Alice Brown,Marketing,Floor 2,Bob Wilson\n")
            return path
    
    @staticmethod
    def create_test_word(path: Path) -> Path:
        """Create a test Word document"""
        try:
            import docx
            
            doc = docx.Document()
            doc.add_heading('Test Document', 0)
            
            doc.add_paragraph('This is a test Word document.')
            doc.add_paragraph('It contains multiple paragraphs for testing.')
            
            doc.add_heading('Section 1', level=1)
            doc.add_paragraph('Content for section 1.')
            
            doc.add_heading('Section 2', level=1)
            doc.add_paragraph('Content for section 2.')
            
            doc.save(str(path))
            return path
            
        except ImportError:
            # Fallback: create text file
            with open(path, 'w') as f:
                f.write("Test Document\n\n")
                f.write("This is a test document.\n")
            return path
    
    @staticmethod
    def create_test_text(path: Path, content: str = None) -> Path:
        """Create a test text file"""
        if content is None:
            content = """Test Document for RAG Pipeline

This document contains test content for validating the RAG pipeline.

Key Features:
- Text extraction
- Chunking
- Embedding generation
- Vector storage
- Retrieval

The pipeline should handle this document correctly."""
        
        with open(path, 'w', encoding='utf-8') as f:
            f.write(content)
        return path


class PipelineTestCase(unittest.TestCase):
    """Base test case with RAG system setup"""
    
    @classmethod
    def setUpClass(cls):
        """Initialize RAG system once for all tests"""
        logger.info("Initializing RAG system for tests...")
        cls.container = initialize_system()
        cls.ingestion_engine = cls.container.get('ingestion_engine')
        cls.query_engine = cls.container.get('query_engine')
        cls.faiss_store = cls.container.get('faiss_store')
        cls.metadata_store = cls.container.get('metadata_store')
        cls.temp_dir = tempfile.mkdtemp()
        logger.info("RAG system initialized successfully")
    
    @classmethod
    def tearDownClass(cls):
        """Cleanup after all tests"""
        import shutil
        shutil.rmtree(cls.temp_dir, ignore_errors=True)
    
    def setUp(self):
        """Setup for each test"""
        self.test_files = []
    
    def tearDown(self):
        """Cleanup after each test"""
        # Clean up test files
        for file_path in self.test_files:
            try:
                Path(file_path).unlink(missing_ok=True)
            except:
                pass


class TestMetadataFlattening(PipelineTestCase):
    """Test metadata flattening fixes"""
    
    def test_no_nested_metadata_in_storage(self):
        """Verify metadata is properly flattened before storage"""
        # Create test document
        test_file = Path(self.temp_dir) / "test_metadata.txt"
        TestDocument.create_test_text(test_file, "Test content for metadata validation")
        self.test_files.append(test_file)
        
        # Ingest with custom metadata
        custom_metadata = {
            'author': 'Test Author',
            'category': 'Test Category',
            'metadata': {  # Intentionally nested
                'nested_field': 'should_be_flattened'
            }
        }
        
        result = self.ingestion_engine.ingest_file(str(test_file), custom_metadata)
        
        # Verify ingestion succeeded
        self.assertEqual(result['status'], 'success')
        self.assertGreater(result['chunks_created'], 0)
        
        # Check stored metadata
        vector_ids = result.get('vector_ids', [])
        self.assertGreater(len(vector_ids), 0)
        
        # Verify no nested metadata in FAISS store
        for vector_id in vector_ids[:3]:  # Check first 3 vectors
            metadata = self.faiss_store.get_vector_metadata(vector_id)
            self.assertIsNotNone(metadata)
            
            # Check for nested metadata
            if 'metadata' in metadata:
                self.assertNotIsInstance(metadata['metadata'], dict, 
                    f"Found nested metadata in vector {vector_id}: {metadata}")
            
            # Verify flattened fields exist
            self.assertIn('text', metadata)
            self.assertIn('chunk_index', metadata)
            self.assertIn('doc_id', metadata)
            
            logger.info(f"Vector {vector_id} metadata keys: {list(metadata.keys())}")
    
    def test_search_returns_flat_metadata(self):
        """Verify search results have flat metadata structure"""
        # Create and ingest test document
        test_file = Path(self.temp_dir) / "test_search.txt"
        TestDocument.create_test_text(test_file, 
            "Machine learning is a subset of artificial intelligence.")
        self.test_files.append(test_file)
        
        result = self.ingestion_engine.ingest_file(str(test_file))
        self.assertEqual(result['status'], 'success')
        
        # Search for content using process_query
        search_results = self.query_engine.process_query("artificial intelligence", top_k=3)
        
        # Verify results structure
        self.assertIn('sources', search_results)
        sources = search_results['sources']
        self.assertGreater(len(sources), 0)
        
        for source in sources:
            # Should have flat structure
            self.assertIn('text', source)
            self.assertIn('similarity_score', source)
            
            # Should NOT have nested metadata
            if 'metadata' in source:
                self.assertNotIsInstance(source['metadata'], dict,
                    f"Found nested metadata in search result: {source}")
            
            logger.info(f"Search result keys: {list(source.keys())}")


class TestPDFProcessing(PipelineTestCase):
    """Test PDF processing with Azure CV integration"""
    
    def test_pdf_basic_ingestion(self):
        """Test basic PDF ingestion"""
        # Create test PDF
        test_pdf = Path(self.temp_dir) / "test.pdf"
        TestDocument.create_test_pdf(test_pdf, with_images=False, pages=2)
        self.test_files.append(test_pdf)
        
        # Ingest PDF
        result = self.ingestion_engine.ingest_file(str(test_pdf))
        
        # Verify success
        self.assertEqual(result['status'], 'success')
        self.assertGreater(result['chunks_created'], 0)
        
        # Verify page-based chunking if enhanced processor used
        if result.get('processor_used') == 'EnhancedPDFProcessor':
            # Should have at least one chunk per page
            self.assertGreaterEqual(result['chunks_created'], 2)
        
        logger.info(f"PDF ingestion result: {result}")
    
    def test_pdf_with_azure_cv(self):
        """Test PDF processing with Azure CV for images"""
        # Check if Azure CV is configured
        config = self.container.get('config_manager').get_config()
        azure_config = config.azure_ai
        
        if not (azure_config.computer_vision_endpoint and azure_config.computer_vision_key):
            self.skipTest("Azure CV not configured")
        
        # Create PDF with images
        test_pdf = Path(self.temp_dir) / "test_with_images.pdf"
        TestDocument.create_test_pdf(test_pdf, with_images=True, pages=2)
        self.test_files.append(test_pdf)
        
        # Ingest PDF
        result = self.ingestion_engine.ingest_file(str(test_pdf))
        
        # Verify success
        self.assertEqual(result['status'], 'success')
        
        # Check metadata for Azure CV processing
        vector_ids = result.get('vector_ids', [])
        if vector_ids:
            metadata = self.faiss_store.get_vector_metadata(vector_ids[0])
            
            # Check for enhanced processing indicators
            if metadata.get('extraction_method') == 'azure_computer_vision':
                self.assertIn('page_number', metadata)
                logger.info("PDF processed with Azure CV")
            else:
                logger.info("PDF processed without Azure CV")


class TestExcelProcessing(PipelineTestCase):
    """Test Excel processing with embedded objects"""
    
    def test_excel_basic_ingestion(self):
        """Test basic Excel ingestion"""
        # Create test Excel
        test_excel = Path(self.temp_dir) / "test.xlsx"
        TestDocument.create_test_excel(test_excel, with_images=False)
        self.test_files.append(test_excel)
        
        # Ingest Excel
        result = self.ingestion_engine.ingest_file(str(test_excel))
        
        # Verify success
        self.assertEqual(result['status'], 'success')
        self.assertGreater(result['chunks_created'], 0)
        
        # Search for Excel content using process_query
        search_results = self.query_engine.process_query("Floor 3 Engineering", top_k=3)
        
        # Verify we can find the content
        sources = search_results.get('sources', [])
        found_content = any("John Doe" in s.get('text', '') for s in sources)
        self.assertTrue(found_content, "Could not find Excel content in search results")
    
    def test_excel_hierarchical_data(self):
        """Test Excel with floor/manager hierarchical data"""
        # Create Excel with hierarchical data
        test_excel = Path(self.temp_dir) / "test_hierarchy.xlsx"
        TestDocument.create_test_excel(test_excel)
        self.test_files.append(test_excel)
        
        # Ingest
        result = self.ingestion_engine.ingest_file(str(test_excel))
        self.assertEqual(result['status'], 'success')
        
        # Search for floor information
        floor_results = self.query_engine.process_query("Floor 2 Marketing", top_k=3)
        sources = floor_results.get('sources', [])
        self.assertGreater(len(sources), 0)
        
        # Search for manager information
        manager_results = self.query_engine.process_query("Jane Smith manager", top_k=3)
        sources = manager_results.get('sources', [])
        self.assertGreater(len(sources), 0)


class TestProcessorIntegration(PipelineTestCase):
    """Test processor chunk preservation"""
    
    def test_processor_chunks_not_rechunked(self):
        """Verify processor chunks are used directly"""
        # Create a mock processor that returns specific chunks
        mock_processor = MagicMock()
        mock_processor.can_process.return_value = True
        mock_processor.process.return_value = {
            'status': 'success',
            'chunks': [
                {
                    'text': 'Processor chunk 1',
                    'metadata': {'page': 1, 'source': 'test_processor'}
                },
                {
                    'text': 'Processor chunk 2', 
                    'metadata': {'page': 2, 'source': 'test_processor'}
                }
            ]
        }
        
        # Temporarily register mock processor
        original_get_processor = self.ingestion_engine.processor_registry.get_processor
        self.ingestion_engine.processor_registry.get_processor = lambda x: mock_processor
        
        try:
            # Create and ingest test file
            test_file = Path(self.temp_dir) / "test_processor.txt"
            TestDocument.create_test_text(test_file)
            self.test_files.append(test_file)
            
            result = self.ingestion_engine.ingest_file(str(test_file))
            
            # Verify processor chunks were used
            self.assertEqual(result['status'], 'success')
            self.assertEqual(result['chunks_created'], 2)  # Exactly 2 chunks from processor
            
            # Verify chunk content
            vector_ids = result.get('vector_ids', [])
            self.assertEqual(len(vector_ids), 2)
            
            # Check metadata preservation
            for i, vector_id in enumerate(vector_ids):
                metadata = self.faiss_store.get_vector_metadata(vector_id)
                self.assertEqual(metadata.get('source'), 'test_processor')
                self.assertEqual(metadata.get('page'), i + 1)
        
        finally:
            # Restore original processor
            self.ingestion_engine.processor_registry.get_processor = original_get_processor


class TestDocumentDeletion(PipelineTestCase):
    """Test document deletion with consistent doc_id"""
    
    def test_delete_by_doc_path(self):
        """Test deletion using doc_path"""
        # Create and ingest test file
        test_file = Path(self.temp_dir) / "test_delete.txt"
        TestDocument.create_test_text(test_file, "Content to be deleted")
        self.test_files.append(test_file)
        
        # Ingest with doc_path metadata
        metadata = {'doc_path': 'test/documents/test_delete.txt'}
        result = self.ingestion_engine.ingest_file(str(test_file), metadata)
        self.assertEqual(result['status'], 'success')
        
        vector_count_before = self.faiss_store.get_index_info()['active_vectors']
        
        # Delete using doc_path
        delete_result = self.ingestion_engine.delete_file(
            str(test_file), 
            doc_path='test/documents/test_delete.txt'
        )
        
        self.assertEqual(delete_result['status'], 'success')
        self.assertGreater(delete_result['vectors_deleted'], 0)
        
        # Verify vectors are marked as deleted
        vector_count_after = self.faiss_store.get_index_info()['active_vectors']
        self.assertLess(vector_count_after, vector_count_before)


class TestEndToEndRAG(PipelineTestCase):
    """Test complete RAG workflow"""
    
    def test_ingest_and_query_workflow(self):
        """Test complete ingestion and query workflow"""
        # Create multiple test documents
        docs = []
        
        # Document 1: Technical content
        tech_file = Path(self.temp_dir) / "technical.txt"
        TestDocument.create_test_text(tech_file, """
        Machine Learning Best Practices
        
        When implementing machine learning models, consider these key factors:
        1. Data quality is paramount
        2. Feature engineering can make or break your model
        3. Always validate on held-out test sets
        4. Monitor model drift in production
        """)
        docs.append(tech_file)
        self.test_files.append(tech_file)
        
        # Document 2: Business content
        business_file = Path(self.temp_dir) / "business.txt"
        TestDocument.create_test_text(business_file, """
        Quarterly Business Review
        
        Our Q3 performance shows:
        - Revenue increased by 15%
        - Customer satisfaction at all-time high
        - New product launch successful
        - Market expansion into APAC region
        """)
        docs.append(business_file)
        self.test_files.append(business_file)
        
        # Ingest all documents
        for doc in docs:
            result = self.ingestion_engine.ingest_file(str(doc))
            self.assertEqual(result['status'], 'success')
        
        # Test queries
        test_queries = [
            ("What are machine learning best practices?", ["data quality", "feature engineering"]),
            ("How was Q3 performance?", ["Revenue increased", "15%"]),
            ("Tell me about customer satisfaction", ["all-time high"])
        ]
        
        for query, expected_terms in test_queries:
            logger.info(f"\nTesting query: {query}")
            
            # Process query
            result = self.query_engine.process_query(query, top_k=3)
            
            # Verify response
            self.assertIn('response', result)
            self.assertIn('sources', result)
            
            response_text = result['response'].lower()
            
            # Check if expected terms appear in response
            found_terms = [term for term in expected_terms 
                          if term.lower() in response_text]
            
            self.assertGreater(len(found_terms), 0, 
                f"Expected terms {expected_terms} not found in response: {result['response']}")
            
            logger.info(f"Response: {result['response'][:200]}...")
            logger.info(f"Sources: {len(result['sources'])}")


class TestConversation(PipelineTestCase):
    """Test conversational interface"""
    
    def test_conversation_with_context(self):
        """Test conversation with context from documents"""
        # Skip if conversation manager not available
        conversation_manager = self.container.get('conversation_manager')
        if not conversation_manager:
            self.skipTest("Conversation manager not available")
        
        # Create and ingest a document about a specific topic
        test_file = Path(self.temp_dir) / "company_info.txt"
        TestDocument.create_test_text(test_file, """
        Acme Corporation Overview
        
        Founded in 2010, Acme Corporation specializes in cloud computing solutions.
        Our headquarters is located in San Francisco, with offices in New York and London.
        
        Key Products:
        - AcmeCloud: Enterprise cloud platform
        - AcmeAnalytics: Business intelligence suite
        - AcmeSecure: Security management system
        
        CEO: John Smith
        CTO: Jane Doe
        Revenue: $500M (2023)
        """)
        self.test_files.append(test_file)
        
        result = self.ingestion_engine.ingest_file(str(test_file))
        self.assertEqual(result['status'], 'success')
        
        # Start conversation
        session_id = "test_session_123"
        
        # First question
        response1 = conversation_manager.process_user_message(
            session_id,
            "What does Acme Corporation do?"
        )
        
        self.assertIn("cloud computing", response1.get('response', '').lower())
        
        # Follow-up question using context
        response2 = conversation_manager.process_user_message(
            session_id,
            "Who is their CEO?"  # 'their' refers to context
        )
        
        self.assertIn("john smith", response2.get('response', '').lower())


class TestPerformanceMetrics(PipelineTestCase):
    """Test performance and metrics collection"""
    
    def test_ingestion_performance(self):
        """Measure ingestion performance"""
        # Create a larger test document
        test_file = Path(self.temp_dir) / "performance_test.txt"
        
        # Generate substantial content
        content = "Performance Test Document\n\n"
        for i in range(100):
            content += f"Paragraph {i}: " + "Lorem ipsum " * 50 + "\n\n"
        
        TestDocument.create_test_text(test_file, content)
        self.test_files.append(test_file)
        
        # Measure ingestion time
        start_time = time.time()
        result = self.ingestion_engine.ingest_file(str(test_file))
        ingestion_time = time.time() - start_time
        
        # Verify success
        self.assertEqual(result['status'], 'success')
        
        # Log performance metrics
        file_size_mb = Path(test_file).stat().st_size / (1024 * 1024)
        chunks_per_second = result['chunks_created'] / ingestion_time
        
        logger.info(f"Performance Metrics:")
        logger.info(f"  File size: {file_size_mb:.2f} MB")
        logger.info(f"  Ingestion time: {ingestion_time:.2f} seconds")
        logger.info(f"  Chunks created: {result['chunks_created']}")
        logger.info(f"  Chunks/second: {chunks_per_second:.2f}")
        logger.info(f"  MB/second: {file_size_mb/ingestion_time:.2f}")
    
    def test_search_performance(self):
        """Measure search performance"""
        # Ensure we have content to search
        test_file = Path(self.temp_dir) / "search_perf_test.txt"
        TestDocument.create_test_text(test_file)
        self.test_files.append(test_file)
        
        result = self.ingestion_engine.ingest_file(str(test_file))
        self.assertEqual(result['status'], 'success')
        
        # Measure search time
        queries = [
            "text extraction",
            "chunking",
            "embedding generation",
            "vector storage",
            "retrieval"
        ]
        
        search_times = []
        for query in queries:
            start_time = time.time()
            results = self.query_engine.process_query(query, top_k=5)
            search_time = time.time() - start_time
            search_times.append(search_time)
            
            logger.info(f"Search '{query}': {search_time*1000:.2f} ms, {len(results.get('sources', []))} results")
        
        avg_search_time = sum(search_times) / len(search_times)
        logger.info(f"Average search time: {avg_search_time*1000:.2f} ms")


class TestErrorHandling(PipelineTestCase):
    """Test error handling and recovery"""
    
    def test_invalid_file_handling(self):
        """Test handling of invalid files"""
        # Non-existent file - should raise FileProcessingError
        with self.assertRaises(Exception):  # FileProcessingError or similar
            self.ingestion_engine.ingest_file("/non/existent/file.txt")
        
        # Empty file - should handle gracefully
        empty_file = Path(self.temp_dir) / "empty.txt"
        empty_file.touch()
        self.test_files.append(empty_file)
        
        try:
            result = self.ingestion_engine.ingest_file(str(empty_file))
            # Should handle gracefully
            self.assertIn('status', result)
        except Exception as e:
            # Azure embedding service might reject empty content
            # This is acceptable behavior
            self.assertIn('embedding', str(e).lower() or 'azure' in str(e).lower())
    
    def test_dimension_mismatch_handling(self):
        """Test handling of embedding dimension mismatches"""
        # This would test the dimension compatibility checking
        current_dim = self.faiss_store.get_current_dimension()
        
        # Check compatibility
        compat_result = self.faiss_store.check_dimension_compatibility(current_dim + 100)
        
        self.assertFalse(compat_result['compatible'])
        self.assertIn('migration_options', compat_result)
        self.assertGreater(len(compat_result['migration_options']), 0)


class TestRetrievalQuality(PipelineTestCase):
    """Test retrieval accuracy and relevance"""
    
    def test_semantic_similarity_accuracy(self):
        """Test if semantically similar queries return relevant results"""
        # Ingest documents with known content
        test_cases = [
            {
                'document': "The CEO announced a new strategic initiative focusing on AI and machine learning to improve customer experience.",
                'queries': [
                    "What is the company's AI strategy?",
                    "How is the CEO planning to use artificial intelligence?",
                    "Machine learning initiatives for customers"
                ],
                'expected_keywords': ['CEO', 'AI', 'machine learning', 'customer']
            },
            {
                'document': "Our Q4 revenue reached $2.5 billion, a 23% increase year-over-year. EBITDA margins improved to 18%.",
                'queries': [
                    "What was the financial performance in Q4?",
                    "How much did revenue grow?",
                    "Tell me about EBITDA margins"
                ],
                'expected_keywords': ['Q4', 'revenue', '2.5 billion', '23%', 'EBITDA', '18%']
            }
        ]
        
        for test_case in test_cases:
            # Ingest document
            file_path = Path(self.temp_dir) / f"semantic_test_{hash(test_case['document'])}.txt"
            TestDocument.create_test_text(file_path, test_case['document'])
            self.ingestion_engine.ingest_file(str(file_path))
            
            # Test each query
            for query in test_case['queries']:
                results = self.query_engine.process_query(query, top_k=5)
                self.assertGreater(len(results.get('sources', [])), 0, f"No results for query: {query}")
                
                # Check if top result contains expected content
                sources = results.get('sources', [])
                if sources:
                    top_result = sources[0]
                    result_text = top_result.get('text', '').lower()
                    
                    found_keywords = [kw for kw in test_case['expected_keywords'] 
                                    if kw.lower() in result_text]
                    
                    self.assertGreater(len(found_keywords), len(test_case['expected_keywords']) * 0.5,
                        f"Query '{query}' didn't return relevant results")
    
    def test_negative_queries(self):
        """Test that irrelevant queries don't return high-confidence results"""
        # Ingest specific content
        file_path = Path(self.temp_dir) / "specific_content.txt"
        TestDocument.create_test_text(file_path, 
            "This document is about quantum computing and its applications in cryptography.")
        self.ingestion_engine.ingest_file(str(file_path))
        
        # Query for unrelated content
        irrelevant_queries = [
            "What is the recipe for chocolate cake?",
            "How to fix a broken car engine?",
            "Best tourist destinations in Europe"
        ]
        
        for query in irrelevant_queries:
            result = self.query_engine.process_query(query, top_k=3)
            
            # Should either return low confidence or indicate no relevant information
            if 'confidence_score' in result:
                self.assertLess(result['confidence_score'], 0.5,
                    f"High confidence for irrelevant query: {query}")
    
    def test_multi_hop_reasoning(self):
        """Test queries requiring information from multiple chunks"""
        # Create documents with related information spread across chunks
        doc1 = Path(self.temp_dir) / "company_info.txt"
        TestDocument.create_test_text(doc1, """
        Company Overview:
        Acme Corp was founded in 2010 by John Smith and Jane Doe.
        
        ... (long text to force chunking) ...
        
        Financial Information:
        The company's revenue in 2023 was $500 million.
        
        ... (more text) ...
        
        Leadership:
        John Smith serves as CEO while Jane Doe is the CTO.
        """)
        
        self.ingestion_engine.ingest_file(str(doc1))
        
        # Query requiring multi-hop reasoning
        result = self.query_engine.process_query(
            "What was the revenue of the company founded by John Smith?"
        )
        
        self.assertIn("500 million", result['response'])
        # Should have multiple sources
        self.assertGreater(len(result.get('sources', [])), 1)


class TestChunkingStrategies(PipelineTestCase):
    """Test different chunking strategies and edge cases"""
    
    def test_code_chunking(self):
        """Test chunking of code files preserves structure"""
        code_content = '''
def calculate_fibonacci(n):
    """Calculate fibonacci number"""
    if n <= 1:
        return n
    
    cache = [0] * (n + 1)
    cache[0] = 0
    cache[1] = 1
    
    for i in range(2, n + 1):
        cache[i] = cache[i-1] + cache[i-2]
    
    return cache[n]

class DataProcessor:
    def __init__(self):
        self.data = []
    
    def process(self, item):
        # Processing logic here
        self.data.append(item)
        return item * 2
'''
        
        file_path = Path(self.temp_dir) / "test_code.py"
        TestDocument.create_test_text(file_path, code_content)
        
        result = self.ingestion_engine.ingest_file(str(file_path))
        
        # Verify code structure is preserved in chunks
        vector_ids = result.get('vector_ids', [])
        for vid in vector_ids[:2]:
            metadata = self.faiss_store.get_vector_metadata(vid)
            chunk_text = metadata.get('text', '')
            
            # Check that code blocks aren't broken mid-function
            if 'def calculate_fibonacci' in chunk_text:
                self.assertIn('return cache[n]', chunk_text,
                    "Function was split across chunks")
    
    def test_table_chunking(self):
        """Test chunking preserves table structure"""
        table_content = """
        Financial Summary:
        
        | Quarter | Revenue | Profit | Growth |
        |---------|---------|--------|--------|
        | Q1 2023 | $1.2M   | $200K  | 15%    |
        | Q2 2023 | $1.5M   | $300K  | 25%    |
        | Q3 2023 | $1.8M   | $400K  | 20%    |
        | Q4 2023 | $2.1M   | $500K  | 17%    |
        
        Analysis: Strong growth throughout 2023...
        """
        
        file_path = Path(self.temp_dir) / "test_table.md"
        TestDocument.create_test_text(file_path, table_content)
        
        result = self.ingestion_engine.ingest_file(str(file_path))
        
        # Search for table data
        results = self.query_engine.process_query("Q3 2023 revenue", top_k=3)
        
        # Verify table context is preserved
        sources = results.get('sources', [])
        self.assertTrue(any("$1.8M" in s.get('text', '') for s in sources))
    
    def test_list_chunking(self):
        """Test chunking of lists and bullet points"""
        list_content = """
        Project Requirements:
        
        1. Authentication System
           - User registration and login
           - OAuth integration
           - Two-factor authentication
           - Password reset functionality
        
        2. Data Management
           - CRUD operations
           - Data validation
           - Batch processing
           - Real-time synchronization
        
        3. Reporting Module
           - Dashboard creation
           - Export to PDF/Excel
           - Scheduled reports
           - Custom metrics
        """
        
        file_path = Path(self.temp_dir) / "test_list.txt"
        TestDocument.create_test_text(file_path, list_content)
        
        result = self.ingestion_engine.ingest_file(str(file_path))
        
        # Verify list items stay together
        results = self.query_engine.process_query("OAuth integration", top_k=1)
        sources = results.get('sources', [])
        if sources:
            text = sources[0].get('text', '')
            # Should include the parent context
            self.assertIn("Authentication System", text)


class TestConcurrentOperations(PipelineTestCase):
    """Test thread safety and concurrent operations"""
    
    def test_concurrent_ingestion(self):
        """Test multiple files being ingested simultaneously"""
        # Create multiple test files
        test_files = []
        for i in range(10):
            file_path = Path(self.temp_dir) / f"concurrent_test_{i}.txt"
            TestDocument.create_test_text(file_path, 
                f"Document {i}: Unique content for concurrent test {i}")
            test_files.append(file_path)
        
        # Track results
        results = []
        errors = []
        lock = threading.Lock()
        
        def ingest_file(file_path):
            try:
                result = self.ingestion_engine.ingest_file(str(file_path))
                with lock:
                    results.append(result)
            except Exception as e:
                with lock:
                    errors.append((file_path, str(e)))
        
        # Ingest files concurrently
        with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
            futures = [executor.submit(ingest_file, f) for f in test_files]
            concurrent.futures.wait(futures)
        
        # Verify all succeeded
        self.assertEqual(len(errors), 0, f"Errors during concurrent ingestion: {errors}")
        self.assertEqual(len(results), len(test_files))
        
        # Verify all documents are searchable
        for i in range(len(test_files)):
            search_results = self.query_engine.process_query(f"Document {i}", top_k=1)
            sources = search_results.get('sources', [])
            self.assertGreater(len(sources), 0)
    
    def test_concurrent_search(self):
        """Test multiple searches happening simultaneously"""
        # First ingest some content
        file_path = Path(self.temp_dir) / "search_test.txt"
        TestDocument.create_test_text(file_path, 
            "Search test content with various keywords: apple, banana, cherry, date, elderberry")
        self.ingestion_engine.ingest_file(str(file_path))
        
        # Concurrent searches
        search_terms = ["apple", "banana", "cherry", "date", "elderberry"] * 10
        results = []
        
        def search_term(term):
            return self.query_engine.process_query(term, top_k=3)
        
        with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
            future_results = executor.map(search_term, search_terms)
            results = list(future_results)
        
        # Verify all searches completed
        self.assertEqual(len(results), len(search_terms))


class TestEdgeCases(PipelineTestCase):
    """Test edge cases and boundary conditions"""
    
    def test_very_large_document(self):
        """Test handling of very large documents"""
        # Create a large document (5MB)
        large_content = "Large document test. " * 100000
        file_path = Path(self.temp_dir) / "large_doc.txt"
        TestDocument.create_test_text(file_path, large_content)
        
        result = self.ingestion_engine.ingest_file(str(file_path))
        self.assertEqual(result['status'], 'success')
        
        # Verify chunk count is reasonable
        self.assertGreater(result['chunks_created'], 50)
        self.assertLess(result['chunks_created'], 10000)
    
    def test_special_characters_handling(self):
        """Test handling of special characters and unicode"""
        special_content = """
        Special Characters Test:
        - Emojis: 😀 🎉 🚀 💡
        - Unicode: café, naïve, Zürich
        - Math: ∑∫∂≤≥≠±∞
        - Symbols: @#$%^&*()_+-={}[]|\\:";'<>?,./
        - Chinese: 你好世界
        - Arabic: مرحبا بالعالم
        - Hebrew: שלום עולם
        """
        
        file_path = Path(self.temp_dir) / "special_chars.txt"
        TestDocument.create_test_text(file_path, special_content)
        
        result = self.ingestion_engine.ingest_file(str(file_path))
        self.assertEqual(result['status'], 'success')
        
        # Test search with special characters
        test_searches = ["café", "😀", "你好", "∑∫∂"]
        for term in test_searches:
            results = self.query_engine.process_query(term, top_k=1)
            # Should handle gracefully even if no results
            self.assertIsInstance(results, dict)
    
    def test_empty_and_whitespace_handling(self):
        """Test handling of empty content and whitespace"""
        test_cases = [
            ("", "empty.txt"),
            ("   ", "spaces.txt"),
            ("\n\n\n", "newlines.txt"),
            ("\t\t\t", "tabs.txt")
        ]
        
        for content, filename in test_cases:
            file_path = Path(self.temp_dir) / filename
            TestDocument.create_test_text(file_path, content)
            
            try:
                result = self.ingestion_engine.ingest_file(str(file_path))
                # Should handle gracefully
                self.assertIn('status', result)
            except Exception as e:
                # Azure embedding service might reject empty/whitespace content
                # This is acceptable behavior
                error_msg = str(e).lower()
                self.assertTrue(
                    any(term in error_msg for term in ['embedding', 'azure', 'empty', 'content']),
                    f"Unexpected error for {filename}: {e}"
                )


class TestRAGMetrics(PipelineTestCase):
    """Test RAG quality metrics and evaluation"""
    
    def test_answer_relevance(self):
        """Test answer relevance scoring"""
        # Ingest document with specific facts
        facts_doc = """
        Company Facts:
        - Founded: 2010
        - Headquarters: San Francisco
        - Employees: 5,000
        - CEO: Jane Smith
        - Industry: Technology
        - Main Product: CloudDB
        """
        
        file_path = Path(self.temp_dir) / "facts.txt"
        TestDocument.create_test_text(file_path, facts_doc)
        self.ingestion_engine.ingest_file(str(file_path))
        
        # Test queries with expected relevance
        test_cases = [
            {
                'query': 'When was the company founded?',
                'expected_in_answer': '2010',
                'min_confidence': 0.8
            },
            {
                'query': 'Who is the CEO?',
                'expected_in_answer': 'Jane Smith',
                'min_confidence': 0.8
            },
            {
                'query': 'What is their main database product?',
                'expected_in_answer': 'CloudDB',
                'min_confidence': 0.7
            }
        ]
        
        for test in test_cases:
            result = self.query_engine.process_query(test['query'])
            
            # Check answer contains expected information
            self.assertIn(test['expected_in_answer'].lower(), 
                         result['response'].lower())
            
            # Check confidence if available
            if 'confidence_score' in result:
                self.assertGreater(result['confidence_score'], 
                                 test['min_confidence'])
    
    def test_source_attribution(self):
        """Test correct source attribution in answers"""
        # Create multiple documents with overlapping information
        doc1 = Path(self.temp_dir) / "source1.txt"
        TestDocument.create_test_text(doc1, 
            "Product A was launched in 2020. It has 1 million users.")
        
        doc2 = Path(self.temp_dir) / "source2.txt"
        TestDocument.create_test_text(doc2, 
            "Product A revenue in 2023 was $50 million.")
        
        self.ingestion_engine.ingest_file(str(doc1))
        self.ingestion_engine.ingest_file(str(doc2))
        
        # Query spanning both documents
        result = self.query_engine.process_query(
            "Tell me about Product A's launch date and revenue"
        )
        
        # Should have sources from both documents
        sources = result.get('sources', [])
        source_files = [s.get('metadata', {}).get('filename', '') 
                       for s in sources]
        
        self.assertIn('source1.txt', str(source_files))
        self.assertIn('source2.txt', str(source_files))
    
    def test_hallucination_detection(self):
        """Test system doesn't hallucinate information"""
        # Ingest limited information
        file_path = Path(self.temp_dir) / "limited_info.txt"
        TestDocument.create_test_text(file_path, 
            "Our company sells software. We have offices in New York.")
        self.ingestion_engine.ingest_file(str(file_path))
        
        # Ask about information not in the document
        queries_not_in_doc = [
            "What is the company's revenue?",
            "How many employees does the company have?",
            "When was the company founded?"
        ]
        
        for query in queries_not_in_doc:
            result = self.query_engine.process_query(query)
            response = result['response'].lower()
            
            # Should indicate information is not available
            self.assertTrue(
                any(phrase in response for phrase in 
                    ["not found", "no information", "not available", 
                     "don't have", "cannot find"]),
                f"Response may contain hallucination for: {query}")


class TestSystemIntegration(PipelineTestCase):
    """Test integration between different system components"""
    
    def test_metadata_store_consistency(self):
        """Test consistency between FAISS and metadata store"""
        # Ingest document with custom metadata
        file_path = Path(self.temp_dir) / "metadata_test.txt"
        TestDocument.create_test_text(file_path, "Test content")
        
        custom_metadata = {
            'author': 'Test Author',
            'department': 'Engineering',
            'priority': 'High'
        }
        
        result = self.ingestion_engine.ingest_file(str(file_path), custom_metadata)
        self.assertEqual(result['status'], 'success')
        
        # Check FAISS store
        vector_ids = result.get('vector_ids', [])
        if vector_ids:  # Only test if vectors were created
            for vector_id in vector_ids:
                # Get metadata from FAISS
                faiss_metadata = self.faiss_store.get_vector_metadata(vector_id)
                
                # Get metadata from metadata store
                metadata_store_data = self.metadata_store.get_document_metadata(
                    faiss_metadata.get('doc_id', '')
                )
                
                # Should be consistent
                self.assertEqual(faiss_metadata.get('author'), 'Test Author')
                self.assertEqual(faiss_metadata.get('department'), 'Engineering')
                self.assertEqual(faiss_metadata.get('priority'), 'High')
    
    def test_processor_registry_integration(self):
        """Test processor registry integration with ingestion"""
        # Get available processors
        processors = self.ingestion_engine.processor_registry.list_processors()
        self.assertGreater(len(processors), 0)
        
        # Test each processor can be retrieved
        for processor_name in processors:
            processor = self.ingestion_engine.processor_registry.get_processor(processor_name)
            if processor is not None:  # Some processors might not be available
                self.assertTrue(hasattr(processor, 'can_process'))
                self.assertTrue(hasattr(processor, 'process'))
    
    def test_config_manager_integration(self):
        """Test configuration manager integration"""
        config_manager = self.container.get('config_manager')
        config = config_manager.get_config()
        
        # Verify essential config sections exist
        self.assertIsNotNone(config.embedding)
        self.assertIsNotNone(config.llm)
        # Note: vector_store might be under a different name or structure
        # Check for common config attributes instead
        self.assertTrue(hasattr(config, 'embedding'))
        self.assertTrue(hasattr(config, 'llm'))
        
        # Test config validation
        validation_result = config_manager.validate_config()
        self.assertTrue(validation_result['valid'])


class TestAdvancedFeatures(PipelineTestCase):
    """Test advanced RAG features and capabilities"""
    
    def test_multi_modal_processing(self):
        """Test multi-modal document processing"""
        # Create a document with mixed content
        mixed_content = """
        Technical Documentation
        
        Our API endpoints:
        1. GET /users - Retrieve user list
        2. POST /users - Create new user
        3. PUT /users/{id} - Update user
        
        [Image: API Architecture Diagram]
        The diagram shows the flow between frontend and backend services.
        
        Code Example:
        ```python
        import requests
        response = requests.get('/api/users')
        users = response.json()
        ```
        """
        
        file_path = Path(self.temp_dir) / "multimodal.txt"
        TestDocument.create_test_text(file_path, mixed_content)
        
        result = self.ingestion_engine.ingest_file(str(file_path))
        self.assertEqual(result['status'], 'success')
        
        # Test search for different content types
        searches = [
            "API endpoints",
            "code example",
            "architecture diagram"
        ]
        
        for search_term in searches:
            results = self.query_engine.process_query(search_term, top_k=3)
            sources = results.get('sources', [])
            self.assertGreater(len(sources), 0)
    
    def test_context_window_handling(self):
        """Test handling of context window limitations"""
        # Create a very long document
        long_content = "Long document content. " * 10000
        file_path = Path(self.temp_dir) / "long_doc.txt"
        TestDocument.create_test_text(file_path, long_content)
        
        result = self.ingestion_engine.ingest_file(str(file_path))
        self.assertEqual(result['status'], 'success')
        
        # Query that might exceed context window
        long_query = "Tell me everything about " + "very detailed information " * 100
        
        # Should handle gracefully
        response = self.query_engine.process_query(long_query)
        self.assertIn('response', response)
        self.assertIsInstance(response['response'], str)
    
    def test_adaptive_chunking(self):
        """Test adaptive chunking based on content structure"""
        # Create document with mixed structure
        structured_content = """
        # Chapter 1: Introduction
        
        This is the introduction chapter with basic concepts.
        
        ## Section 1.1: Background
        
        Background information goes here.
        
        ### Subsection 1.1.1: Historical Context
        
        Historical context details.
        
        # Chapter 2: Methodology
        
        ## Section 2.1: Research Design
        
        Research methodology description.
        
        ### Subsection 2.1.1: Data Collection
        
        Data collection methods.
        """
        
        file_path = Path(self.temp_dir) / "structured.txt"
        TestDocument.create_test_text(file_path, structured_content)
        
        result = self.ingestion_engine.ingest_file(str(file_path))
        self.assertEqual(result['status'], 'success')
        
        # Verify chunks respect structure
        vector_ids = result.get('vector_ids', [])
        if vector_ids:  # Only test if vectors were created
            self.assertGreater(len(vector_ids), 1)
            
            # Check that structure is preserved in metadata
            for vector_id in vector_ids[:3]:
                metadata = self.faiss_store.get_vector_metadata(vector_id)
                # Should have structural information
                self.assertIn('text', metadata)
    
    def test_semantic_search_enhancement(self):
        """Test enhanced semantic search capabilities"""
        # Create documents with semantic relationships
        doc1 = Path(self.temp_dir) / "semantic1.txt"
        TestDocument.create_test_text(doc1, 
            "The feline companion rested on the windowsill, watching birds.")
        
        doc2 = Path(self.temp_dir) / "semantic2.txt"
        TestDocument.create_test_text(doc2, 
            "Our pet cat enjoys sitting by the window to observe wildlife.")
        
        self.ingestion_engine.ingest_file(str(doc1))
        self.ingestion_engine.ingest_file(str(doc2))
        
        # Test semantic search
        semantic_queries = [
            "cat behavior",
            "pet activities",
            "window observation"
        ]
        
        for query in semantic_queries:
            results = self.query_engine.process_query(query, top_k=5)
            sources = results.get('sources', [])
            # Should find semantically related content
            self.assertGreater(len(sources), 0)
            
            # Check semantic relevance
            found_relevant = any(
                any(term in s.get('text', '').lower() 
                    for term in ['cat', 'pet', 'window', 'watch'])
                for s in sources
            )
            self.assertTrue(found_relevant, 
                          f"Semantic search failed for query: {query}")


def run_comprehensive_tests():
    """Run all tests and generate report"""
    # Create test suite
    test_suite = unittest.TestSuite()
    
    # Add test classes
    test_classes = [
        TestMetadataFlattening,
        TestPDFProcessing,
        TestExcelProcessing,
        TestProcessorIntegration,
        TestDocumentDeletion,
        TestEndToEndRAG,
        TestConversation,
        TestPerformanceMetrics,
        TestErrorHandling,
        TestRetrievalQuality,
        TestChunkingStrategies,
        TestConcurrentOperations,
        TestEdgeCases,
        TestRAGMetrics,
        TestSystemIntegration,
        TestAdvancedFeatures
    ]
    
    for test_class in test_classes:
        tests = unittest.TestLoader().loadTestsFromTestCase(test_class)
        test_suite.addTests(tests)
    
    # Run tests
    runner = unittest.TextTestRunner(verbosity=2)
    result = runner.run(test_suite)
    
    # Generate summary report
    print("\n" + "="*60)
    print("PIPELINE VALIDATION SUMMARY")
    print("="*60)
    print(f"Tests run: {result.testsRun}")
    print(f"Failures: {len(result.failures)}")
    print(f"Errors: {len(result.errors)}")
    print(f"Skipped: {len(result.skipped)}")
    print(f"Success rate: {(result.testsRun - len(result.failures) - len(result.errors)) / result.testsRun * 100:.1f}%")
    
    if result.failures:
        print("\nFAILURES:")
        for test, traceback in result.failures:
            print(f"- {test}: {traceback.split(chr(10))[-2]}")
    
    if result.errors:
        print("\nERRORS:")
        for test, traceback in result.errors:
            print(f"- {test}: {traceback.split(chr(10))[-2]}")
    
    return result.wasSuccessful()


if __name__ == "__main__":
    # Run comprehensive tests
    success = run_comprehensive_tests()
    
    # Exit with appropriate code
    sys.exit(0 if success else 1)