#!/usr/bin/env python3
"""
Script to create a test PDF document for RAG system ingestion testing
Requires: pip install reportlab
"""

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    import os
except ImportError:
    print("❌ ReportLab not installed. Installing...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab"])
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

def create_test_pdf():
    """Create a comprehensive test PDF document"""
    
    # Create the PDF file
    filename = "employee_handbook.pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter,
                          rightMargin=72, leftMargin=72,
                          topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    story = []
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        spaceAfter=12,
        textColor=colors.darkblue
    )
    
    # Title
    story.append(Paragraph("EMPLOYEE HANDBOOK", title_style))
    story.append(Paragraph("Technology Solutions Inc.", styles['Heading3']))
    story.append(Paragraph("Effective Date: January 1, 2025", styles['Normal']))
    story.append(Spacer(1, 20))
    
    # Table of Contents
    story.append(Paragraph("TABLE OF CONTENTS", heading_style))
    toc_data = [
        ['Section', 'Page'],
        ['1. Welcome Message', '2'],
        ['2. Company Overview', '2'],
        ['3. Employment Policies', '3'],
        ['4. Benefits and Compensation', '4'],
        ['5. Code of Conduct', '5'],
        ['6. IT and Security Policies', '6'],
        ['7. Emergency Procedures', '7'],
        ['8. Contact Information', '8']
    ]
    
    toc_table = Table(toc_data, colWidths=[4*inch, 1*inch])
    toc_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(toc_table)
    story.append(Spacer(1, 20))
    
    # Section 1: Welcome Message
    story.append(Paragraph("1. WELCOME MESSAGE", heading_style))
    welcome_text = """
    Welcome to Technology Solutions Inc.! We are excited to have you join our team of innovative 
    professionals dedicated to delivering cutting-edge technology solutions to our clients worldwide.
    
    This handbook contains important information about our company policies, procedures, benefits, 
    and expectations. Please read it carefully and keep it as a reference throughout your employment.
    
    Our success depends on the contributions of each team member. We encourage open communication, 
    continuous learning, and collaborative problem-solving. Together, we will continue to grow and 
    achieve excellence in everything we do.
    """
    story.append(Paragraph(welcome_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section 2: Company Overview
    story.append(Paragraph("2. COMPANY OVERVIEW", heading_style))
    
    story.append(Paragraph("Mission Statement", styles['Heading4']))
    mission_text = """
    To provide innovative technology solutions that empower businesses to achieve their goals 
    while maintaining the highest standards of quality, security, and customer service.
    """
    story.append(Paragraph(mission_text, styles['Normal']))
    
    story.append(Paragraph("Core Values", styles['Heading4']))
    values_data = [
        ['Value', 'Description'],
        ['Innovation', 'We embrace new technologies and creative solutions'],
        ['Integrity', 'We conduct business with honesty and transparency'],
        ['Excellence', 'We strive for the highest quality in all our work'],
        ['Collaboration', 'We work together to achieve common goals'],
        ['Customer Focus', 'We prioritize our customers\' success and satisfaction']
    ]
    
    values_table = Table(values_data, colWidths=[1.5*inch, 3.5*inch])
    values_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(values_table)
    story.append(Spacer(1, 12))
    
    # Section 3: Employment Policies
    story.append(Paragraph("3. EMPLOYMENT POLICIES", heading_style))
    
    story.append(Paragraph("Work Schedule", styles['Heading4']))
    schedule_text = """
    Standard business hours are Monday through Friday, 9:00 AM to 5:00 PM EST. We offer flexible 
    work arrangements including hybrid remote work options. Employees may work from home up to 
    2 days per week with manager approval.
    """
    story.append(Paragraph(schedule_text, styles['Normal']))
    
    story.append(Paragraph("Attendance Policy", styles['Heading4']))
    attendance_text = """
    Regular attendance is essential for business operations. If you cannot report to work, 
    notify your supervisor as soon as possible, preferably before your scheduled start time. 
    Excessive absenteeism may result in disciplinary action.
    """
    story.append(Paragraph(attendance_text, styles['Normal']))
    
    story.append(Paragraph("Performance Reviews", styles['Heading4']))
    performance_text = """
    Performance evaluations are conducted annually, with quarterly check-ins throughout the year. 
    Reviews assess job performance, goal achievement, and professional development opportunities. 
    Employees are encouraged to provide feedback and discuss career aspirations during these sessions.
    """
    story.append(Paragraph(performance_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section 4: Benefits and Compensation
    story.append(Paragraph("4. BENEFITS AND COMPENSATION", heading_style))
    
    benefits_text = """
    We offer a comprehensive benefits package designed to support your health, financial security, 
    and work-life balance:
    """
    story.append(Paragraph(benefits_text, styles['Normal']))
    
    benefits_data = [
        ['Benefit', 'Details'],
        ['Health Insurance', 'Medical, dental, and vision coverage with company contribution'],
        ['Retirement Plan', '401(k) with 4% company match after 90 days'],
        ['Paid Time Off', '15 days vacation, 10 sick days, 12 holidays annually'],
        ['Professional Development', '$2,000 annual budget for training and conferences'],
        ['Life Insurance', 'Company-paid life insurance equal to annual salary'],
        ['Flexible Spending', 'FSA accounts for healthcare and dependent care expenses']
    ]
    
    benefits_table = Table(benefits_data, colWidths=[2*inch, 3*inch])
    benefits_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkgreen),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightgreen),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'TOP')
    ]))
    
    story.append(benefits_table)
    story.append(Spacer(1, 12))
    
    # Section 5: Code of Conduct
    story.append(Paragraph("5. CODE OF CONDUCT", heading_style))
    
    conduct_text = """
    All employees are expected to maintain the highest standards of professional conduct:
    
    • Treat all colleagues, clients, and partners with respect and dignity
    • Maintain confidentiality of proprietary and client information
    • Avoid conflicts of interest and disclose potential conflicts to management
    • Comply with all applicable laws, regulations, and company policies
    • Report any suspected violations of this code to HR or management
    • Use company resources responsibly and for business purposes only
    
    Violations of this code may result in disciplinary action, up to and including termination.
    """
    story.append(Paragraph(conduct_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section 6: IT and Security Policies
    story.append(Paragraph("6. IT AND SECURITY POLICIES", heading_style))
    
    story.append(Paragraph("Password Requirements", styles['Heading4']))
    password_text = """
    • Minimum 12 characters with uppercase, lowercase, numbers, and special characters
    • Change passwords every 90 days
    • Do not reuse the last 12 passwords
    • Enable multi-factor authentication on all business accounts
    """
    story.append(Paragraph(password_text, styles['Normal']))
    
    story.append(Paragraph("Data Protection", styles['Heading4']))
    data_text = """
    • Encrypt all devices containing company or client data
    • Use approved cloud storage services only
    • Report data breaches immediately to the security team
    • Follow clean desk policy - secure all sensitive materials
    """
    story.append(Paragraph(data_text, styles['Normal']))
    
    story.append(Paragraph("Acceptable Use", styles['Heading4']))
    use_text = """
    Company IT resources are provided for business use. Limited personal use is permitted 
    but should not interfere with work responsibilities. Prohibited activities include:
    accessing inappropriate content, installing unauthorized software, or using resources 
    for personal business ventures.
    """
    story.append(Paragraph(use_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section 7: Emergency Procedures
    story.append(Paragraph("7. EMERGENCY PROCEDURES", heading_style))
    
    emergency_text = """
    In case of emergency, follow these procedures:
    
    Fire Emergency:
    1. Evacuate immediately using nearest exit
    2. Do not use elevators
    3. Proceed to designated assembly area in parking lot
    4. Wait for all-clear from emergency personnel
    
    Medical Emergency:
    1. Call 911 immediately
    2. Notify building security at ext. 9999
    3. Provide first aid if trained and safe to do so
    4. Do not move injured person unless in immediate danger
    
    Security Incident:
    1. Report to security immediately at ext. 9999
    2. Do not confront suspicious individuals
    3. Secure your work area if safe to do so
    4. Follow instructions from security personnel
    """
    story.append(Paragraph(emergency_text, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # Section 8: Contact Information
    story.append(Paragraph("8. CONTACT INFORMATION", heading_style))
    
    contact_data = [
        ['Department', 'Phone', 'Email'],
        ['Human Resources', 'ext. 4200', 'hr@techsolutions.com'],
        ['IT Help Desk', 'ext. 4357', 'helpdesk@techsolutions.com'],
        ['Security', 'ext. 9999', 'security@techsolutions.com'],
        ['Facilities', 'ext. 4500', 'facilities@techsolutions.com'],
        ['Emergency Line', '911', 'N/A']
    ]
    
    contact_table = Table(contact_data, colWidths=[2*inch, 1.5*inch, 2*inch])
    contact_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkorange),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.lightyellow),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(contact_table)
    story.append(Spacer(1, 20))
    
    # Footer
    footer_text = """
    This handbook is effective as of January 1, 2025, and supersedes all previous versions. 
    The company reserves the right to modify these policies at any time with appropriate notice.
    
    For questions about this handbook, please contact Human Resources.
    
    Document Version: 2.1
    Last Updated: 2025-06-07
    """
    story.append(Paragraph(footer_text, styles['Normal']))
    
    # Build PDF
    doc.build(story)
    print(f"✅ Created test PDF: {filename}")
    return filename

def create_simple_docx():
    """Create a simple DOCX test document"""
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        print("❌ python-docx not installed. Installing...")
        import subprocess
        import sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
        from docx.shared import Inches
    
    # Create document
    doc = Document()
    
    # Title
    title = doc.add_heading('IT Security Guidelines', 0)
    
    # Metadata
    doc.add_paragraph('Document Type: Microsoft Word (.docx)')
    doc.add_paragraph('Created: 2025-06-07')
    doc.add_paragraph('Purpose: Test document for RAG system ingestion')
    
    # Introduction
    doc.add_heading('Introduction', level=1)
    intro_text = """
    This document outlines the essential IT security guidelines that all employees must follow 
    to protect our organization's digital assets and maintain compliance with industry standards.
    """
    doc.add_paragraph(intro_text)
    
    # Password Security
    doc.add_heading('Password Security', level=1)
    doc.add_paragraph('Strong passwords are the first line of defense against cyber threats.')
    
    doc.add_heading('Password Requirements:', level=2)
    password_list = doc.add_paragraph()
    password_list.add_run('• Minimum 12 characters in length\n')
    password_list.add_run('• Include uppercase and lowercase letters\n')
    password_list.add_run('• Include at least one number\n')
    password_list.add_run('• Include at least one special character\n')
    password_list.add_run('• Avoid dictionary words or personal information\n')
    password_list.add_run('• Change passwords every 90 days')
    
    # Multi-Factor Authentication
    doc.add_heading('Multi-Factor Authentication (MFA)', level=1)
    mfa_text = """
    MFA is required for all business applications and systems. This adds an extra layer of 
    security by requiring a second form of verification beyond your password.
    """
    doc.add_paragraph(mfa_text)
    
    doc.add_heading('Supported MFA Methods:', level=2)
    mfa_list = doc.add_paragraph()
    mfa_list.add_run('• Authenticator apps (Google Authenticator, Microsoft Authenticator)\n')
    mfa_list.add_run('• SMS text messages\n')
    mfa_list.add_run('• Hardware security keys\n')
    mfa_list.add_run('• Biometric authentication (fingerprint, face recognition)')
    
    # Email Security
    doc.add_heading('Email Security', level=1)
    email_text = """
    Email is a common vector for cyber attacks. Follow these guidelines to protect yourself 
    and the organization from email-based threats.
    """
    doc.add_paragraph(email_text)
    
    doc.add_heading('Best Practices:', level=2)
    email_list = doc.add_paragraph()
    email_list.add_run('• Verify sender identity before clicking links or downloading attachments\n')
    email_list.add_run('• Be suspicious of urgent requests for sensitive information\n')
    email_list.add_run('• Report phishing attempts to the security team\n')
    email_list.add_run('• Use encrypted email for sensitive communications\n')
    email_list.add_run('• Keep software and email clients updated')
    
    # Data Protection
    doc.add_heading('Data Protection', level=1)
    data_text = """
    Protecting sensitive data is everyone's responsibility. This includes customer information, 
    financial data, intellectual property, and personal employee information.
    """
    doc.add_paragraph(data_text)
    
    doc.add_heading('Data Classification Levels:', level=2)
    
    # Add a table for data classification
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Classification'
    hdr_cells[1].text = 'Description'
    hdr_cells[2].text = 'Examples'
    
    classifications = [
        ('Public', 'Information that can be freely shared', 'Marketing materials, press releases'),
        ('Internal', 'Information for internal use only', 'Employee directories, internal policies'),
        ('Confidential', 'Sensitive business information', 'Financial reports, customer lists'),
        ('Restricted', 'Highly sensitive information', 'Personal data, trade secrets')
    ]
    
    for classification, description, examples in classifications:
        row_cells = table.add_row().cells
        row_cells[0].text = classification
        row_cells[1].text = description
        row_cells[2].text = examples
    
    # Incident Reporting
    doc.add_heading('Security Incident Reporting', level=1)
    incident_text = """
    If you suspect a security incident, report it immediately to the security team. 
    Quick reporting can help minimize damage and prevent further compromise.
    """
    doc.add_paragraph(incident_text)
    
    doc.add_heading('What to Report:', level=2)
    incident_list = doc.add_paragraph()
    incident_list.add_run('• Suspected malware infections\n')
    incident_list.add_run('• Unauthorized access attempts\n')
    incident_list.add_run('• Lost or stolen devices\n')
    incident_list.add_run('• Suspicious emails or phone calls\n')
    incident_list.add_run('• Data breaches or unauthorized data access')
    
    # Contact Information
    doc.add_heading('Emergency Contacts', level=1)
    contact_text = """
    Security Team: security@company.com or ext. 9999 (24/7)
    IT Help Desk: helpdesk@company.com or ext. 4357
    HR Department: hr@company.com or ext. 4200
    """
    doc.add_paragraph(contact_text)
    
    # Footer
    doc.add_paragraph('\n' + '='*50)
    footer_text = """
    This document is confidential and proprietary. Distribution is restricted to authorized 
    personnel only. For questions or clarifications, contact the IT Security team.
    
    Document Version: 1.3
    Last Updated: 2025-06-07
    Next Review Date: 2025-12-07
    """
    doc.add_paragraph(footer_text)
    
    # Save document
    filename = "it_security_guidelines.docx"
    doc.save(filename)
    print(f"✅ Created test DOCX: {filename}")
    return filename

if __name__ == "__main__":
    print("🔧 Creating test documents for RAG system ingestion...")
    
    # Create PDF
    try:
        pdf_file = create_test_pdf()
    except Exception as e:
        print(f"❌ Failed to create PDF: {e}")
    
    # Create DOCX
    try:
        docx_file = create_simple_docx()
    except Exception as e:
        print(f"❌ Failed to create DOCX: {e}")
    
    print("\n✅ Test document creation completed!")
    print("\nCreated files:")
    print("- employee_handbook.pdf (comprehensive employee handbook)")
    print("- it_security_guidelines.docx (IT security policies)")
    print("- sample_knowledge_base.txt (technical knowledge base)")
    print("- api_documentation.md (API documentation with rich formatting)")
    print("\nThese documents can be used to test:")
    print("• Different file format ingestion (PDF, DOCX, TXT, MD)")
    print("• Various content types (policies, procedures, technical docs)")
    print("• Rich formatting preservation (tables, lists, headings)")
    print("• Metadata extraction and search capabilities") 